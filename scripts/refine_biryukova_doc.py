# -*- coding: utf-8 -*-
from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Бирюкова.docx"
SCREEN_DIR = ROOT / "tmp" / "real_screenshots"


def set_run_font(run, size: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run, 12)
    return new_para


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def normalize_body_paragraph(paragraph, left: bool = False) -> None:
    paragraph.style = "No Spacing"
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.JUSTIFY


def make_text_paragraph(paragraph, text: str, left: bool = False) -> None:
    clear_paragraph(paragraph)
    normalize_body_paragraph(paragraph, left=left)
    run = paragraph.add_run(text)
    set_run_font(run, 12)


def add_after(anchor, texts: list[str], left: bool = False):
    current = anchor
    for text in texts:
        current = insert_paragraph_after(current, style="No Spacing")
        normalize_body_paragraph(current, left=left)
        run = current.add_run(text)
        set_run_font(run, 12)
    return current


def find_paragraph(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def replace_image_blob(doc: Document, shape, image_bytes: bytes) -> None:
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    rid = blip.embed
    image_part = doc.part.related_parts[rid]
    image_part._blob = image_bytes


def prepare_image(src: Path, crop_height: int | None = None, canvas_size: tuple[int, int] = (1600, 980)) -> bytes:
    with Image.open(src) as original:
        image = original.convert("RGB")
        if crop_height is not None:
            crop_height = min(crop_height, image.height)
            image = image.crop((0, 0, image.width, crop_height))

        canvas = Image.new("RGB", canvas_size, "#ffffff")
        max_w = canvas_size[0] - 60
        max_h = canvas_size[1] - 60
        ratio = min(max_w / image.width, max_h / image.height)
        new_size = (int(image.width * ratio), int(image.height * ratio))
        fitted = image.resize(new_size, Image.Resampling.LANCZOS)
        x = (canvas_size[0] - new_size[0]) // 2
        y = (canvas_size[1] - new_size[1]) // 2
        canvas.paste(fitted, (x, y))
        out = BytesIO()
        canvas.save(out, format="PNG")
        return out.getvalue()


def actual_tree_lines() -> list[str]:
    nbsp = "\u00A0"
    return [
        "app/",
        f"{nbsp*2}\u251c\u2500\u2500 Console/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Commands/",
        f"{nbsp*2}\u2502{nbsp*3}\u2514\u2500\u2500 Kernel.php",
        f"{nbsp*2}\u251c\u2500\u2500 Helpers/",
        f"{nbsp*2}\u251c\u2500\u2500 Http/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Controllers/",
        f"{nbsp*2}\u2502{nbsp*3}\u2502{nbsp*3}\u2514\u2500\u2500 Auth/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Middleware/",
        f"{nbsp*2}\u2502{nbsp*3}\u2514\u2500\u2500 Requests/",
        f"{nbsp*2}\u251c\u2500\u2500 Models/",
        f"{nbsp*2}\u2514\u2500\u2500 Providers/",
        "database/",
        f"{nbsp*2}\u251c\u2500\u2500 factories/",
        f"{nbsp*2}\u251c\u2500\u2500 migrations/",
        f"{nbsp*2}\u2514\u2500\u2500 seeders/",
        "public/",
        f"{nbsp*2}\u251c\u2500\u2500 build/",
        f"{nbsp*2}\u2514\u2500\u2500 qr_codes/",
        "resources/",
        f"{nbsp*2}\u251c\u2500\u2500 css/",
        f"{nbsp*2}\u251c\u2500\u2500 js/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Components/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Lang/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Layouts/",
        f"{nbsp*2}\u2502{nbsp*3}\u251c\u2500\u2500 Pages/",
        f"{nbsp*2}\u2502{nbsp*3}\u2502{nbsp*3}\u2514\u2500\u2500 Admin/",
        f"{nbsp*2}\u2502{nbsp*3}\u2514\u2500\u2500 utils/",
        f"{nbsp*2}\u251c\u2500\u2500 lang/",
        f"{nbsp*2}\u2514\u2500\u2500 views/",
        "routes/",
        f"{nbsp*2}\u251c\u2500\u2500 api.php",
        f"{nbsp*2}\u251c\u2500\u2500 auth.php",
        f"{nbsp*2}\u251c\u2500\u2500 health.php",
        f"{nbsp*2}\u2514\u2500\u2500 web.php",
        "scripts/",
        "storage/",
        "tests/",
    ]


def refine_document() -> None:
    doc = Document(str(DOC_PATH))

    # Replace only actual interface screenshots, keeping diagram slots untouched.
    screenshot_map = {
        3: prepare_image(SCREEN_DIR / "qr_generator.png", crop_height=1250),
        4: prepare_image(SCREEN_DIR / "qr_history.png", crop_height=1250),
        5: prepare_image(SCREEN_DIR / "qr_analytics.png", crop_height=1350),
        6: prepare_image(SCREEN_DIR / "qr_scanner.png", crop_height=1450),
        7: prepare_image(SCREEN_DIR / "admin_dashboard.png", crop_height=1250),
        8: prepare_image(SCREEN_DIR / "admin_qr_codes.png", crop_height=1350),
    }
    for idx, blob in screenshot_map.items():
        replace_image_blob(doc, doc.inline_shapes[idx], blob)

    # Replace figure 2.3 with text tree of the real repository.
    tree_anchor = find_paragraph(doc, "Инфраструктурный слой размещен в верхнеуровневых файлах `composer.json`, `package.json`, `Dockerfile`, `docker-compose.yml`, а маршруты собраны в каталоге `routes`. В сумме структура проекта демонстрирует хорошую инженерную дисциплину: backend, frontend, схема БД и конфигурация развертывания не смешаны между собой и образуют понятный каркас приложения.")
    insert_after = tree_anchor
    for line in actual_tree_lines():
        insert_after = insert_paragraph_after(insert_after, style="No Spacing")
        insert_after.paragraph_format.first_line_indent = Cm(0)
        insert_after.paragraph_format.line_spacing = 1.5
        insert_after.paragraph_format.space_after = Pt(0)
        insert_after.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = insert_after.add_run(line)
        set_run_font(run, 12)

    fig23_caption = find_paragraph(doc, "Рисунок 2.3 – Структура каталогов проекта")
    fig23_image_paragraph = None
    for p in doc.paragraphs:
        if "<w:drawing" in p._p.xml and p._p.getnext() is fig23_caption._p:
            fig23_image_paragraph = p
            break
    if fig23_image_paragraph is not None:
        delete_paragraph(fig23_image_paragraph)
    delete_paragraph(fig23_caption)

    # Expand sections with more detailed text.
    add_after(
        find_paragraph(doc, "С инженерной точки зрения backend проекта выглядит достаточно зрелым для учебной работы: контроллеры отделены по ролям, модели компактны и понятны, а административный модуль использует общие возможности схемы БД без дублирования CRUD-кода. Это создает хороший фундамент для дальнейшего вынесения части логики в сервисный слой при росте требований."),
        [
            "Если рассматривать маршрутизацию подробнее, то backend фактически разделен на четыре независимых контура. Первый контур образуют публичные страницы `/` и `/scan`, доступные без авторизации. Второй контур объединяет защищенные пользовательские маршруты `/generate`, `/history`, `/profile`, `/plans`, `/feedback` и `/qr/{id}/analytics`. Третий контур представлен административным префиксом `/admin`, где один контроллер обслуживает несколько таблиц. Четвертый контур формирует публичный redirect-маршрут `/r/{slug}`, который принципиально отличается от остальных тем, что должен обрабатываться быстро, без лишних зависимостей и без ожидания авторизации.",
            "Отдельного комментария заслуживает метод `index` контроллера QR-кодов. Он не просто возвращает список записей, а одновременно применяет поиск, фильтрацию, сортировку, подсчет количества сканирований и тарифные ограничения на показ динамических кодов. Именно здесь хорошо видно, как backend учитывает бизнес-логику приложения: пользователь бесплатного плана не должен видеть dynamic QR в истории, а пользователь платного тарифа получает полный реестр созданных записей. На выходе формируется не сырой набор моделей, а готовый payload для интерфейса истории.",
            "Не менее показателен `AdminController`, построенный как универсальная административная прослойка над схемой базы данных. Контроллер получает список таблиц, исключает служебные сущности, динамически определяет метаданные колонок и уже на этой основе строит формы и таблицы. Такой подход нельзя назвать идеальным для крупного enterprise-продукта, однако для учебно-прикладного сервиса он оправдан: разработчик быстро получает обозримый инструмент управления данными и одновременно сохраняет одну точку развития административного функционала.",
        ],
    )
    add_after(
        find_paragraph(doc, "Административный frontend вынесен в отдельные страницы `Admin/Index.vue`, `Admin/TableIndex.vue` и `Admin/TableForm.vue`. Такое решение поддерживает четкое различие между пользовательским и служебным интерфейсом, но при этом не разрывает единое визуальное и навигационное пространство приложения."),
        [
            "Важно, что frontend не дублирует серверную логику, а работает как тонкий клиент над Inertia-props. Например, `QrHistory.vue` не вычисляет право на доступ к dynamic QR и не считает статистику самостоятельно: страница только отображает уже подготовленные backend-данные, управляет локальным состоянием фильтров и инициирует переходы. Благодаря этому приложение избегает рассинхронизации между тем, что видно пользователю, и тем, что реально разрешено сервером.",
            "В генераторе QR-кодов особенно заметна роль клиентского слоя как интерфейса подготовки данных. Страница собирает поля для разных типов содержимого, преобразует их к финальной строке, показывает превью на canvas и позволяет пользователю оценить результат еще до сохранения. При этом окончательное решение о допустимости dynamic-режима остается на сервере. Такой баланс между client-side UX и server-side control можно считать удачным: пользователь получает удобную интерактивность, а ключевые ограничения не покидают защищенный контур backend.",
            "Страницы аналитики и сканера также хорошо демонстрируют разделение обязанностей. `QrAnalytics.vue` отвечает за визуализацию уже готового набора событий, тогда как `QrScanner.vue` работает с браузерным API камеры, списком устройств и библиотекой распознавания. Иными словами, frontend закрывает те задачи, которые естественно выполнять в браузере, а все, что связано с данными, безопасностью и целостностью состояния, остается в Laravel.",
        ],
    )
    add_after(
        find_paragraph(doc, "Таким образом, генерация и хранение в проекте реализованы корректно: создание кода, файловое сохранение, доступ к истории и публичное использование одного и того же объекта опираются на единый набор данных и не требуют дублирующих сущностей."),
        [
            "Если проследить сценарий до уровня полей и параметров запроса, то видно, что сервер сохраняет не только строковое значение кода, но и контекст его построения. Поля `size`, `color_dark` и `color_light` делают запись воспроизводимой: пользователь может не просто увидеть, что когда-то создавал QR-код, а фактически восстановить его графические характеристики. Поле `payload` выполняет другую роль: оно сохраняет структурированные исходные данные и тем самым упрощает интерпретацию записи в интерфейсе истории.",
            "Файловое хранение в `public/qr_codes` тоже является важным проектным решением. С одной стороны, оно обеспечивает простой доступ к PNG-изображениям без отдельного object storage. С другой стороны, такое решение накладывает требование на аккуратную работу с путями, правами каталога и удалением файлов при очистке записей. В контроллере удаления эти детали уже учтены: приложение сначала пытается убрать физический файл и только затем удаляет запись из базы данных. Это снижает риск накопления «сиротских» изображений.",
            "С практической точки зрения такой сценарий хорошо подходит для небольшого и среднего сервиса. Он не требует отдельного медиа-сервиса, CDN или асинхронного пайплайна генерации. В то же время архитектура не закрывает путь к дальнейшему развитию: при росте нагрузки генерацию изображений и хранение файлов можно вынести в отдельный слой, не меняя пользовательский интерфейс и базовую предметную модель.",
        ],
    )
    add_after(
        find_paragraph(doc, "В результате аналитика, журнал сканирований и поддержка разных типов данных образуют связанный функциональный блок, который выгодно отличает рассматриваемый сервис от простых генераторов изображений."),
        [
            "Рассматривая redirect-логику более детально, важно подчеркнуть, что для dynamic QR ключевым становится не факт выдачи изображения, а качество последующей трассировки переходов. Каждый запрос к `/r/{slug}` превращается в событие, которое связывает конкретный код, время обращения и характеристики клиента. Именно из этих первичных событий затем строятся графики на странице аналитики. Если бы приложение записывало только агрегированное значение `scan_count`, оно не смогло бы показать распределение по странам, устройствам и браузерам.",
            "Текущая реализация аналитики достаточно скромна по сравнению с промышленными маркетинговыми платформами, но в рамках ВКР это скорее достоинство, чем недостаток. Решение остается прозрачным: можно буквально проследить путь данных от redirect-контроллера до таблицы `qr_scans`, затем до запроса на страницу аналитики и наконец до визуального графика на клиенте. Такая прослеживаемость делает систему удобной для сопровождения, верификации и дальнейших доработок.",
            "Поддержка нескольких типов QR-кодов также влияет на аналитическую ценность сервиса. Код, содержащий URL, телефон, письмо или географическую координату, имеет разные сценарии использования и разные прикладные метрики. Даже при единой схеме хранения поле `type` позволяет различать эти сценарии и в дальнейшем строить более точные отчеты: например, сравнивать активность URL-кодов маркетинговой кампании и служебных гео-меток для внутренней навигации.",
        ],
    )
    add_after(
        find_paragraph(doc, "Следовательно, безопасность в проекте реализована не как отдельный декларативный раздел, а как набор конкретных защитных механизмов, встроенных в маршруты, контроллеры и модель доступа."),
        [
            "Следует отметить и то, что безопасность здесь тесно связана с корректностью пользовательского опыта. Если не проверять владельца QR-кода при открытии аналитики, то проблема будет не только в разглашении данных, но и в разрушении доверия к системе как личному рабочему кабинету. Аналогично, неконтролируемый redirect опасен не только с точки зрения security, но и с точки зрения репутации сервиса: пользователь должен понимать, что система перенаправляет его именно туда, куда был настроен конкретный код.",
            "Для административного контура важна еще одна деталь: универсальный CRUD облегчает поддержку, но одновременно требует аккуратного блок-листа таблиц и строгой проверки прав. В проекте эта мера уже реализована через исключение системных таблиц и middleware `admin`. Благодаря этому оператор получает функциональный инструмент управления данными, но не может случайно вмешаться в инфраструктурные служебные сущности, которые не предназначены для редактирования через UI.",
        ],
    )
    add_after(
        find_paragraph(doc, "В целом административный модуль подтверждает, что проект ориентирован не только на конечного пользователя, но и на эксплуатацию: данные можно контролировать, агрегаты — анализировать, а записи — оперативно редактировать через веб-интерфейс."),
        [
            "С эксплуатационной точки зрения особенно полезно то, что dashboard опирается на те же таблицы, которые затем доступны в CRUD-режиме. Администратор может сначала увидеть аномалию на сводном экране, а затем сразу перейти к деталям в таблицах `qr_scans` или `qr_codes`. Такая связка между обзором и детализацией обычно появляется только в более зрелых продуктах, поэтому ее наличие повышает прикладную ценность рассматриваемого решения.",
            "Экран таблицы `qr_codes` также выполняет методическую функцию в контексте ВКР. Через него хорошо видны реальные поля предметной модели: тип, содержимое, payload, принадлежность пользователю, признак dynamic-режима и slug. По сути, административный UI становится визуальным подтверждением того, что архитектурные и модельные решения главы 2 действительно реализованы в работающем приложении, а не остались на уровне декларативного описания.",
            "Наконец, административная панель полезна и как средство миграции к более сложным эксплуатационным сценариям. Даже если в будущем проект получит отдельные отчеты, роли операторов и журналы аудита, уже текущая реализация дает основу для такого развития: данные централизованы, таблицы обозримы, а показатели собираются в одном месте. Это делает переход к следующему уровню зрелости скорее эволюционным, чем революционным.",
        ],
    )
    add_after(
        find_paragraph(doc, "Инфраструктурный слой размещен в верхнеуровневых файлах `composer.json`, `package.json`, `Dockerfile`, `docker-compose.yml`, а маршруты собраны в каталоге `routes`. В сумме структура проекта демонстрирует хорошую инженерную дисциплину: backend, frontend, схема БД и конфигурация развертывания не смешаны между собой и образуют понятный каркас приложения."),
        [
            "Дополнительным достоинством такой структуры является локализация ответственности. Разработчик, работающий с пользовательским интерфейсом, практически сразу попадает в `resources/js/Pages` и связанные layouts, тогда как backend-логика сосредоточена в `app/Http` и моделях. Это уменьшает время на поиск нужной точки изменения и снижает риск того, что правка интерфейса случайно затронет серверный слой или наоборот.",
        ],
    )
    add_after(
        find_paragraph(doc, "В качестве хранилища проект использует реляционную модель данных, совместимую с SQLite и MySQL. Такой выбор естественен для сервиса, где есть пользователи, тарифы, QR-коды и события сканирования, связанные отношениями один-ко-многим. Реляционная СУБД позволяет сохранять целостность, строить выборки по аналитике и поддерживать административный CRUD без дополнительного слоя хранилищ."),
        [
            "Для учебного проекта это особенно важно: стек остается современным, но при этом не перегруженным. Каждая технология имеет четкую функцию, а не используется ради модного названия. Laravel отвечает за доменную и HTTP-логику, Vue и Inertia закрывают интерактивность интерфейсов, Chart.js решает локальную задачу визуализации, а `simple-qrcode` устраняет необходимость вручную реализовывать алгоритм построения графического кода.",
        ],
    )

    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    refine_document()
