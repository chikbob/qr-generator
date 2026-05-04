from __future__ import annotations

import json
import math
import re
import textwrap
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "Исиченко.docx"
OUTPUT_PATH = ROOT / "Бирюкова.docx"
ASSETS_DIR = ROOT / "tmp" / "biryukova_assets"

PAGE_BG = "#ffffff"
TEXT = "#1f2933"
MUTED = "#52606d"
LINE = "#d9e2ec"
ACCENT = "#bd6592"
ACCENT_2 = "#e095bc"
ACCENT_3 = "#f6d6e4"
GOOD = "#2f855a"
BLUE = "#2b6cb0"


def font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    if mono:
        path = "/System/Library/Fonts/Menlo.ttc"
    elif bold:
        path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
        if not Path(path).exists():
            path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    else:
        path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    return ImageFont.truetype(path, size=size)


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, font_obj: ImageFont.FreeTypeFont) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font_obj: ImageFont.FreeTypeFont,
    fill: str = TEXT,
    max_width: int | None = None,
    line_gap: int = 6,
) -> int:
    x, y = xy
    lines = wrap(draw, text, max_width, font_obj) if max_width else text.splitlines()
    for line in lines:
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += font_obj.size + line_gap
    return y


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str, fill: str = "#f8fafc", outline: str = LINE) -> None:
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    draw.text((x1 + 18, y1 + 14), title, font=font(24, bold=True), fill=ACCENT)
    draw_multiline(draw, (x1 + 18, y1 + 52), body, font(18), fill=TEXT, max_width=(x2 - x1 - 36))


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = MUTED, width: int = 4) -> None:
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def canvas(size: tuple[int, int], title: str, subtitle: str | None = None) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, PAGE_BG)
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((24, 24, size[0] - 24, size[1] - 24), radius=28, outline=LINE, width=3, fill="#fffdfd")
    draw.text((48, 42), title, font=font(34, bold=True), fill=TEXT)
    if subtitle:
        draw.text((48, 86), subtitle, font=font(18), fill=MUTED)
    return image, draw


def save(image: Image.Image, name: str) -> Path:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSETS_DIR / name
    image.save(path)
    return path


def draw_architecture() -> Path:
    image, draw = canvas((1700, 1000), "Архитектурная диаграмма веб-сервиса QR-кодов", "Клиентский контур, backend, сервисы хранения и аналитики")
    rounded_box(draw, (70, 180, 360, 360), "Пользователь", "Генерация QR-кода\nПросмотр истории\nАналитика переходов", fill="#fff4f8", outline=ACCENT_3)
    rounded_box(draw, (70, 420, 360, 640), "Администратор", "Управление записями\nСводные показатели\nCRUD по сущностям", fill="#fdf2f8", outline=ACCENT_3)
    rounded_box(draw, (470, 150, 780, 320), "Frontend", "Vue 3 + Inertia.js\nСтраницы генератора, истории, аналитики,\nсканера и административной панели")
    rounded_box(draw, (470, 380, 780, 610), "Laravel backend", "Маршрутизация, контроллеры,\nвалидация, авторизация,\nгенерация изображений QR и редиректы")
    rounded_box(draw, (470, 670, 780, 890), "Сервисный контур", "Подготовка payload, нормализация ссылок,\nучет сканирований, формирование статистики")
    rounded_box(draw, (910, 120, 1230, 300), "Хранилище данных", "users\nplans\nqr_codes\nqr_scans\nfeedback")
    rounded_box(draw, (910, 360, 1230, 540), "Файловое хранилище", "public/qr_codes\nPNG-изображения\nдинамические ссылки")
    rounded_box(draw, (910, 610, 1230, 820), "Внешние сервисы", "ipinfo.io для геолокации\nбраузерные API камеры\nпочтовые и платежные заглушки")
    rounded_box(draw, (1330, 240, 1610, 450), "Публичный redirect-контур", "Маршрут /r/{slug}\nфиксация события сканирования\nперенаправление по целевому URI", fill="#f0fff4", outline="#bee3c7")
    rounded_box(draw, (1330, 560, 1610, 790), "Модуль аналитики", "Агрегация по странам,\nбраузерам, устройствам,\nвременным интервалам", fill="#ebf8ff", outline="#bee3f8")
    for s, e in [
        ((360, 270), (470, 235)),
        ((360, 530), (470, 505)),
        ((780, 235), (910, 210)),
        ((780, 480), (910, 450)),
        ((780, 760), (910, 720)),
        ((1230, 450), (1330, 345)),
        ((1230, 720), (1330, 675)),
        ((470, 505), (360, 270)),
    ]:
        arrow(draw, s, e, fill=ACCENT if s[0] < e[0] else BLUE)
    return save(image, "architecture.png")


def draw_er() -> Path:
    image, draw = canvas((1800, 1200), "ER-диаграмма модели данных", "Целевой набор сущностей веб-сервиса генерации и учета QR-кодов")
    boxes = {
        "users": (90, 150, 390, 360, "users", "id PK\nplan_id FK\nname\nemail\npassword\nis_admin\ncreated_at"),
        "plans": (90, 430, 390, 620, "plans", "id PK\nname\nprice\nis_default"),
        "projects": (510, 120, 820, 340, "projects", "id PK\nuser_id FK\nname\ndescription\nis_archived"),
        "qr_categories": (510, 410, 820, 620, "qr_categories", "id PK\nproject_id FK\nname\nslug\nsort_order"),
        "qr_codes": (930, 100, 1280, 420, "qr_codes", "id PK\nuser_id FK\nproject_id FK\ncategory_id FK\ntype\ncontent\npayload\nimage_path\nis_dynamic\nslug\nscan_count"),
        "qr_scans": (930, 500, 1280, 840, "qr_scans", "id PK\nqr_code_id FK\nip\ncountry\ncity\nbrowser\ndevice\nreferer\ncreated_at"),
        "scan_statistics": (1410, 140, 1710, 360, "scan_statistics", "id PK\nqr_code_id FK\nperiod_date\nscans_total\nunique_devices\nunique_countries"),
        "feedback": (1410, 470, 1710, 680, "feedback", "id PK\nuser_id FK\nsubject\ncategory\nstatus\npriority"),
    }
    for x1, y1, x2, y2, title, body in boxes.values():
        rounded_box(draw, (x1, y1, x2, y2), title, body)
    relations = [
        ((390, 255), (510, 220), "1:N"),
        ((390, 520), (510, 520), "1:N"),
        ((820, 220), (930, 220), "1:N"),
        ((820, 520), (930, 280), "1:N"),
        ((390, 255), (930, 180), "1:N"),
        ((1280, 260), (1410, 250), "1:N"),
        ((1280, 680), (1410, 250), "N:1"),
        ((1280, 680), (1410, 575), "N:1"),
    ]
    for start, end, label in relations:
        arrow(draw, start, end, fill=BLUE)
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 18
        draw.text((mx, my), label, font=font(18, bold=True), fill=ACCENT)
    return save(image, "er.png")


def draw_use_case() -> Path:
    image, draw = canvas((1700, 1000), "UML Use-Case Diagram", "Основные пользовательские и административные сценарии")
    draw.rectangle((320, 140, 1370, 860), outline=LINE, width=3, fill="#fcfcfd")
    draw.text((750, 160), "Веб-сервис QR-кодов", font=font(28, bold=True), fill=TEXT)
    actors = {
        "Гость": (110, 260),
        "Пользователь": (110, 520),
        "Администратор": (1510, 440),
    }
    for label, (x, y) in actors.items():
        draw.ellipse((x - 24, y - 70, x + 24, y - 22), outline=TEXT, width=3)
        draw.line((x, y - 22, x, y + 42), fill=TEXT, width=3)
        draw.line((x - 36, y, x + 36, y), fill=TEXT, width=3)
        draw.line((x, y + 42, x - 30, y + 90), fill=TEXT, width=3)
        draw.line((x, y + 42, x + 30, y + 90), fill=TEXT, width=3)
        draw.text((x - 52, y + 100), label, font=font(20, bold=True), fill=TEXT)
    use_cases = [
        ((520, 260, 860, 320), "Просмотр главной страницы"),
        ((520, 360, 860, 420), "Сканирование QR-кода"),
        ((520, 470, 860, 530), "Генерация QR-кода"),
        ((520, 570, 860, 630), "Управление историей QR"),
        ((520, 670, 860, 730), "Просмотр аналитики"),
        ((960, 320, 1280, 380), "Управление тарифом"),
        ((960, 430, 1280, 490), "Работа с обратной связью"),
        ((960, 540, 1280, 600), "Администрирование таблиц"),
        ((960, 650, 1280, 710), "Просмотр глобальной статистики"),
    ]
    for box, label in use_cases:
        draw.ellipse(box, outline=ACCENT, width=3, fill="#fff7fb")
        x1, y1, x2, y2 = box
        draw_multiline(draw, (x1 + 28, y1 + 14), label, font(18, bold=True), fill=TEXT, max_width=(x2 - x1 - 56), line_gap=3)
    for s, e in [
        ((146, 520), (520, 500)),
        ((146, 520), (520, 600)),
        ((146, 520), (520, 700)),
        ((146, 260), (520, 290)),
        ((146, 260), (520, 390)),
        ((1510, 440), (1280, 570)),
        ((1510, 440), (1280, 680)),
        ((146, 520), (960, 350)),
        ((146, 520), (960, 460)),
    ]:
        draw.line([s, e], fill=MUTED, width=3)
    return save(image, "use_case.png")


def draw_sequence() -> Path:
    image, draw = canvas((1800, 1200), "UML Sequence Diagram", "Сценарий создания динамического QR-кода и последующего сканирования")
    participants = ["Пользователь", "Frontend", "QrCodeController", "qr_codes", "Redirect /r/{slug}", "qr_scans"]
    xs = [160, 450, 760, 1050, 1360, 1630]
    for x, label in zip(xs, participants):
        draw.rounded_rectangle((x - 90, 120, x + 90, 180), radius=16, outline=LINE, width=3, fill="#f8fafc")
        draw_multiline(draw, (x - 72, 136), label, font(18, bold=True), max_width=144, line_gap=2)
        draw.line((x, 180, x, 1080), fill="#cbd5e0", width=3)
    steps = [
        (0, 1, 250, "Ввод параметров и выбор dynamic"),
        (1, 2, 330, "POST /qr с type, content, payload"),
        (2, 3, 410, "INSERT qr_codes"),
        (2, 1, 490, "302 /history + flash"),
        (1, 0, 570, "Отображение карточки QR"),
        (0, 4, 690, "Сканирование /r/{slug}"),
        (4, 5, 770, "INSERT qr_scans"),
        (4, 0, 850, "HTTP redirect на target URI"),
        (0, 1, 960, "Запрос страницы аналитики"),
        (1, 2, 1040, "GET /qr/{id}/analytics"),
        (2, 3, 1120, "SELECT qr_codes + scans_count"),
        (2, 5, 1200, "SELECT qr_scans"),
        (2, 1, 1280, "JSON props для Inertia"),
    ]
    for src, dst, y, label in steps:
        arrow(draw, (xs[src], y), (xs[dst], y), fill=ACCENT if src < dst else BLUE)
        draw.text((min(xs[src], xs[dst]) + 16, y - 28), label, font=font(17), fill=TEXT)
    return save(image, "sequence.png")


def draw_component() -> Path:
    image, draw = canvas((1700, 1000), "Component Diagram", "Ключевые программные компоненты и зависимости проекта")
    items = [
        ((90, 180, 430, 360), "Pages", "Home.vue\nQrGenerator.vue\nQrHistory.vue\nQrAnalytics.vue\nQrScanner.vue"),
        ((90, 430, 430, 660), "Layouts / Components", "AppLayout.vue\nAdminLayout.vue\nHeader/Footer\nAdminHeader"),
        ((520, 170, 860, 390), "HTTP layer", "routes/web.php\nroutes/auth.php\nmiddleware auth/admin\nInertia helpers"),
        ((520, 450, 860, 760), "Controllers", "QrCodeController\nAdminController\nPlanController\nProfileController\nFeedbackController"),
        ((950, 150, 1290, 420), "Domain model", "User\nPlan\nQrCode\nQrScan\nFeedback"),
        ((950, 500, 1290, 760), "Persistence", "migrations\nseeders\nSQLite / MySQL abstraction\nEloquent relations"),
        ((1380, 240, 1660, 420), "Libraries", "simple-qrcode\njenssegers/agent\nsanctum\nchart.js"),
        ((1380, 540, 1660, 770), "Infrastructure", "public/qr_codes\nipinfo.io\nbrowser camera API\nVite build"),
    ]
    for box, title, body in items:
        rounded_box(draw, box, title, body)
    arrows = [
        ((430, 270), (520, 270)),
        ((430, 540), (520, 610)),
        ((860, 300), (950, 280)),
        ((860, 610), (950, 610)),
        ((1290, 280), (1380, 330)),
        ((1290, 620), (1380, 650)),
        ((860, 610), (1380, 650)),
    ]
    for s, e in arrows:
        arrow(draw, s, e, fill=MUTED)
    return save(image, "component.png")


def draw_ui_generator() -> Path:
    image, draw = canvas((1600, 980), "Экран генерации QR-кода", "Фронтенд-интерфейс пользователя")
    draw.rounded_rectangle((70, 130, 1530, 900), radius=26, fill="#fff7fb", outline=ACCENT_3, width=3)
    draw.rounded_rectangle((100, 170, 620, 860), radius=20, fill="#ffffff", outline=LINE, width=2)
    draw.text((130, 210), "Тип QR-кода", font=font(24, bold=True), fill=TEXT)
    fields = [
        ("Выбран тип", "URL / ссылка"),
        ("Целевое значение", "https://qrapp.com/r/demo"),
        ("Размер", "260 px"),
        ("Темный цвет", "#000000"),
        ("Светлый цвет", "#FFFFFF"),
        ("Динамический режим", "Включен"),
    ]
    y = 260
    for label, value in fields:
        draw.text((130, y), label, font=font(18, bold=True), fill=MUTED)
        draw.rounded_rectangle((130, y + 30, 560, y + 90), radius=14, fill="#f8fafc", outline=LINE)
        draw.text((150, y + 48), value, font=font(18), fill=TEXT)
        y += 110
    draw.rounded_rectangle((720, 170, 1480, 860), radius=20, fill="#ffffff", outline=LINE, width=2)
    draw.text((1010, 220), "Превью QR", font=font(28, bold=True), fill=ACCENT)
    for i in range(14):
        for j in range(14):
            if (i * j + i + j) % 3 == 0 or (2 < i < 5 and 2 < j < 5):
                draw.rectangle((910 + i * 32, 320 + j * 32, 936 + i * 32, 346 + j * 32), fill=TEXT)
    draw.rounded_rectangle((900, 700, 1280, 770), radius=18, fill=ACCENT, outline=ACCENT)
    draw.text((1035, 722), "Сохранить", font=font(22, bold=True), fill="#ffffff")
    draw.rounded_rectangle((900, 790, 1280, 850), radius=18, fill="#edf2f7", outline=LINE)
    draw.text((1010, 810), "Скачать PNG", font=font(20, bold=True), fill=TEXT)
    return save(image, "ui_generator.png")


def draw_ui_history() -> Path:
    image, draw = canvas((1600, 980), "Экран истории QR-кодов", "Карточки, фильтрация и быстрые действия")
    draw.rounded_rectangle((70, 130, 1530, 900), radius=26, fill="#f8fafc", outline=LINE, width=3)
    draw.rounded_rectangle((100, 170, 1500, 250), radius=18, fill="#ffffff", outline=LINE)
    filters = ["Поиск по содержимому", "Фильтр: dynamic", "Сортировка: новые сверху"]
    x = 130
    for item in filters:
        draw.rounded_rectangle((x, 190, x + 330, 230), radius=12, fill="#fff7fb", outline=ACCENT_3)
        draw.text((x + 16, 202), item, font=font(17), fill=MUTED)
        x += 360
    cards = [
        ("Dynamic URL", "https://qrapp.com/r/promo-2026", 124, True),
        ("Text", "Инструкция по подключению к сети склада", 52, False),
        ("Geo", "geo:55.7558,37.6176", 38, True),
        ("Email", "mailto:support@qrapp.com", 11, False),
    ]
    top = 310
    lefts = [120, 790]
    idx = 0
    for row in range(2):
        for col in range(2):
            title, content, scans, dynamic = cards[idx]
            x1 = lefts[col]
            y1 = top + row * 250
            draw.rounded_rectangle((x1, y1, x1 + 640, y1 + 190), radius=20, fill="#ffffff", outline=LINE)
            draw.rounded_rectangle((x1 + 20, y1 + 22, x1 + 140, y1 + 142), radius=18, fill="#fff7fb", outline=ACCENT_3)
            for i in range(8):
                for j in range(8):
                    if (i + j) % 2 == 0:
                        draw.rectangle((x1 + 36 + i * 12, y1 + 38 + j * 12, x1 + 44 + i * 12, y1 + 46 + j * 12), fill=TEXT)
            draw.text((x1 + 170, y1 + 28), title, font=font(24, bold=True), fill=TEXT)
            draw_multiline(draw, (x1 + 170, y1 + 62), content, font(17), fill=MUTED, max_width=430)
            tag = "Dynamic" if dynamic else "Static"
            draw.rounded_rectangle((x1 + 170, y1 + 122, x1 + 280, y1 + 160), radius=12, fill="#edfdf3" if dynamic else "#edf2f7", outline="#bee3c7" if dynamic else LINE)
            draw.text((x1 + 194, y1 + 132), tag, font=font(16, bold=True), fill=GOOD if dynamic else TEXT)
            draw.text((x1 + 310, y1 + 132), f"Сканирования: {scans}", font=font(16), fill=TEXT)
            idx += 1
    return save(image, "ui_history.png")


def draw_ui_analytics() -> Path:
    image, draw = canvas((1600, 980), "Экран аналитики QR-кода", "Линейный график активности и детализация сканирований")
    draw.rounded_rectangle((70, 130, 1530, 900), radius=26, fill="#ffffff", outline=LINE, width=3)
    rounded_box(draw, (100, 170, 520, 360), "Карточка QR", "Тип: dynamic URL\nSlug: sale-spring\nВсего сканов: 124\nСоздан: 12.04.2026")
    draw.rounded_rectangle((560, 170, 1490, 520), radius=20, fill="#f8fafc", outline=LINE)
    draw.text((590, 205), "Активность по дням", font=font(24, bold=True), fill=TEXT)
    points = [(620, 430), (720, 390), (820, 360), (920, 300), (1020, 340), (1120, 250), (1220, 270), (1320, 220), (1420, 240)]
    for i in range(len(points) - 1):
        draw.line((points[i], points[i + 1]), fill=ACCENT, width=5)
    for x, y in points:
        draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=ACCENT, outline=ACCENT)
    draw.rounded_rectangle((100, 560, 1490, 860), radius=20, fill="#fff7fb", outline=ACCENT_3)
    headers = ["Дата", "Страна", "Город", "Браузер", "Устройство", "Referer"]
    col_x = [130, 330, 520, 760, 980, 1210]
    for x, header in zip(col_x, headers):
        draw.text((x, 590), header, font=font(18, bold=True), fill=ACCENT)
    rows = [
        ("14.04 11:42", "RU", "Moscow", "Chrome", "Android", "telegram"),
        ("14.04 10:17", "PL", "Warsaw", "Safari", "iPhone", "qrapp.com"),
        ("13.04 19:03", "DE", "Berlin", "Chrome", "Windows", "campaign"),
        ("13.04 15:20", "RU", "Kazan", "Firefox", "Macintosh", "direct"),
    ]
    y = 635
    for row in rows:
        draw.rounded_rectangle((120, y - 10, 1470, y + 34), radius=10, fill="#ffffff", outline=LINE)
        for x, cell in zip(col_x, row):
            draw.text((x, y), cell, font=font(16), fill=TEXT)
        y += 58
    return save(image, "ui_analytics.png")


def draw_ui_scanner() -> Path:
    image, draw = canvas((1600, 980), "Экран сканирования QR-кодов", "Использование камеры устройства и отображение результата")
    draw.rounded_rectangle((80, 150, 980, 860), radius=24, fill="#0f172a", outline="#1e293b", width=3)
    draw.rounded_rectangle((170, 240, 890, 690), radius=24, fill="#111827", outline="#334155", width=4)
    draw.rectangle((360, 310, 700, 650), outline=ACCENT_2, width=6)
    draw.line((360, 310, 430, 310), fill=ACCENT_2, width=10)
    draw.line((360, 310, 360, 380), fill=ACCENT_2, width=10)
    draw.line((700, 310, 630, 310), fill=ACCENT_2, width=10)
    draw.line((700, 310, 700, 380), fill=ACCENT_2, width=10)
    draw.line((360, 650, 430, 650), fill=ACCENT_2, width=10)
    draw.line((360, 650, 360, 580), fill=ACCENT_2, width=10)
    draw.line((700, 650, 630, 650), fill=ACCENT_2, width=10)
    draw.line((700, 650, 700, 580), fill=ACCENT_2, width=10)
    draw.text((350, 710), "Камера активна. Наведите код в рамку", font=font(22, bold=True), fill="#ffffff")
    rounded_box(draw, (1060, 180, 1490, 420), "Результат сканирования", "Распознан URI:\nhttps://qrapp.com/r/sale-spring\nТип: redirect\nИсточник: back camera")
    rounded_box(draw, (1060, 470, 1490, 760), "Доступные действия", "Открыть ссылку\nОчистить результат\nСменить устройство камеры\nОбновить список камер")
    return save(image, "ui_scanner.png")


def draw_ui_admin_dashboard() -> Path:
    image, draw = canvas((1600, 980), "Административная панель: dashboard", "KPI, топ стран и последние сканирования")
    draw.rounded_rectangle((70, 130, 1530, 900), radius=26, fill="#fffafc", outline=ACCENT_3, width=3)
    kpis = [("Всего сканов", "100"), ("Сканов сегодня", "9"), ("Доля dynamic", "61%"), ("Новых пользователей", "4")]
    x = 100
    for label, value in kpis:
        draw.rounded_rectangle((x, 180, x + 320, 310), radius=18, fill="#ffffff", outline=LINE)
        draw.text((x + 24, 208), label, font=font(18, bold=True), fill=MUTED)
        draw.text((x + 24, 244), value, font=font(36, bold=True), fill=ACCENT)
        x += 350
    rounded_box(draw, (100, 360, 430, 780), "Top countries", "RU — 34\nBY — 18\nPL — 12\nDE — 10\nUS — 8")
    draw.rounded_rectangle((470, 360, 1490, 780), radius=20, fill="#ffffff", outline=LINE)
    draw.text((510, 390), "Последние сканирования", font=font(24, bold=True), fill=TEXT)
    headers = ["Дата", "Страна", "Город", "Браузер", "Устройство", "Контент"]
    cols = [510, 680, 840, 980, 1120, 1280]
    for x, h in zip(cols, headers):
        draw.text((x, 440), h, font=font(17, bold=True), fill=ACCENT)
    rows = [
        ("14.04 12:02", "RU", "Moscow", "Chrome", "Android", "https://qrapp..."),
        ("14.04 11:44", "DE", "Berlin", "Safari", "iPhone", "mailto:info..."),
        ("14.04 11:29", "PL", "Warsaw", "Chrome", "Windows", "geo:55.75..."),
        ("14.04 10:57", "RU", "Kazan", "Firefox", "Macintosh", "promo text"),
    ]
    y = 485
    for row in rows:
        draw.rounded_rectangle((500, y - 8, 1460, y + 34), radius=10, fill="#f8fafc", outline=LINE)
        for x, cell in zip(cols, row):
            draw.text((x, y), cell, font=font(15), fill=TEXT)
        y += 64
    return save(image, "ui_admin_dashboard.png")


def draw_ui_admin_table() -> Path:
    image, draw = canvas((1600, 980), "Административная панель: управление сущностями", "Табличный CRUD для записей qr_codes")
    draw.rounded_rectangle((70, 130, 1530, 900), radius=26, fill="#ffffff", outline=LINE, width=3)
    draw.rounded_rectangle((100, 170, 1500, 240), radius=16, fill="#fff7fb", outline=ACCENT_3)
    draw.text((130, 194), "Таблица: qr_codes", font=font(24, bold=True), fill=TEXT)
    draw.rounded_rectangle((1180, 186, 1460, 224), radius=12, fill=ACCENT, outline=ACCENT)
    draw.text((1258, 196), "Новая запись", font=font(18, bold=True), fill="#ffffff")
    headers = ["id", "user_id", "type", "content", "is_dynamic", "slug", "actions"]
    cols = [120, 220, 340, 450, 1020, 1180, 1330]
    for x, h in zip(cols, headers):
        draw.text((x, 290), h, font=font(17, bold=True), fill=ACCENT)
    rows = [
        ("91", "12", "url", "https://qrapp.com/r/promo", "true", "promo-2026", "edit | delete"),
        ("90", "7", "text", "Инструкция по отгрузке", "false", "—", "edit | delete"),
        ("89", "4", "geo", "geo:55.75,37.61", "true", "office-map", "edit | delete"),
        ("88", "4", "email", "mailto:info@qrapp.com", "false", "—", "edit | delete"),
    ]
    y = 340
    for row in rows:
        draw.rounded_rectangle((110, y - 8, 1470, y + 42), radius=10, fill="#f8fafc", outline=LINE)
        for x, cell in zip(cols, row):
            draw.text((x, y), cell, font=font(15), fill=TEXT)
        y += 72
    return save(image, "ui_admin_table.png")


def draw_project_structure() -> Path:
    image, draw = canvas((1600, 980), "Структура проекта", "Фрагмент дерева каталогов и ключевых файлов")
    draw.rounded_rectangle((70, 130, 1530, 900), radius=24, fill="#0f172a", outline="#1e293b", width=3)
    tree = """.
app
  Http
    Controllers
      AdminController.php
      FeedbackController.php
      PlanController.php
      ProfileController.php
      QrCodeController.php
    Middleware
      AdminMiddleware.php
resources
  js
    Pages
      Home.vue
      QrGenerator.vue
      QrHistory.vue
      QrAnalytics.vue
      QrScanner.vue
      Admin
database
  migrations
  seeders
routes
  web.php
  auth.php
  api.php
public
  qr_codes
composer.json
package.json"""
    draw_multiline(draw, (120, 180), tree, font(24, mono=True), fill="#e2e8f0", line_gap=8)
    return save(image, "project_structure.png")


def set_run_font(run, size: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def add_text(doc: Document, text: str, style: str = "No Spacing") -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if style in {"No Spacing", "Normal"}:
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, 14 if level == 1 else 12, bold=True if level == 1 else False)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 12)


def add_picture(doc: Document, path: Path, width_cm: float = 15.8) -> None:
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_text(cell, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, 11, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths and i < len(widths):
            header_cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT)
            if widths and i < len(widths):
                cells[i].width = Cm(widths[i])


def page_break(doc: Document) -> None:
    doc.add_page_break()


def remove_body_after_prefix(doc: Document, keep_body_index: int) -> None:
    body = doc._body._body
    children = list(body)
    for idx, child in enumerate(children):
        tag = child.tag.rsplit("}", 1)[-1]
        if idx > keep_body_index and tag != "sectPr":
            body.remove(child)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def extract_routes() -> list[list[str]]:
    route_files = [ROOT / "routes" / name for name in ["web.php", "auth.php", "api.php", "health.php"]]
    records: list[list[str]] = []
    pattern = re.compile(
        r"Route::(?P<method>get|post|put|patch|delete|resource)\((?P<args>.*?)\)(?P<tail>[^;]*);",
        re.S,
    )
    name_pattern = re.compile(r"->name\('([^']+)'\)")
    only_pattern = re.compile(r"->only\(\[([^\]]+)\]\)")
    for file in route_files:
        text = file.read_text(encoding="utf-8")
        for match in pattern.finditer(text):
            method = match.group("method").upper()
            args = " ".join(match.group("args").split())
            tail = match.group("tail")
            if method == "RESOURCE":
                uri_match = re.match(r"'([^']+)'", args)
                if not uri_match:
                    continue
                uri = uri_match.group(1)
                only = only_pattern.search(tail)
                actions = [a.strip(" '") for a in only.group(1).split(",")] if only else ["index", "store", "show", "update", "destroy"]
                for action, http, suffix in [
                    ("index", "GET", uri),
                    ("store", "POST", uri),
                    ("destroy", "DELETE", f"{uri}/{{id}}"),
                ]:
                    if action in actions:
                        records.append([http, f"/{suffix}".replace("//", "/"), f"{uri}.{action}", "resource action"])
                continue
            uri_match = re.match(r"'([^']+)'|\"([^\"]+)\"", args)
            uri = uri_match.group(1) or uri_match.group(2) if uri_match else "—"
            route_name = name_pattern.search(tail)
            name = route_name.group(1) if route_name else "—"
            if "QrCodeController::class" in args:
                handler = "QrCodeController"
            elif "AdminController::class" in args:
                handler = "AdminController"
            elif "ProfileController::class" in args:
                handler = "ProfileController"
            elif "PlanController::class" in args:
                handler = "PlanController"
            elif "FeedbackController::class" in args:
                handler = "FeedbackController"
            elif "::class" in args:
                handler = "auth controller"
            elif "inertiaWithUser" in args:
                handler = "Inertia page"
            elif "response()->json" in args:
                handler = "health closure"
            else:
                handler = "closure"
            records.append([method, f"/{uri}".replace("//", "/"), name, handler])
    unique: list[list[str]] = []
    seen = set()
    for row in records:
        key = tuple(row)
        if key not in seen:
            seen.add(key)
            unique.append(row)
    return unique


def chunked(seq: list[list[str]], size: int) -> Iterable[list[list[str]]]:
    for i in range(0, len(seq), size):
        yield seq[i : i + size]


def build_assets() -> dict[str, Path]:
    return {
        "1.1": draw_architecture(),
        "2.1": draw_er(),
        "2.2": draw_use_case(),
        "3.1": draw_sequence(),
        "3.2": draw_component(),
        "3.3": draw_ui_generator(),
        "3.4": draw_ui_history(),
        "3.5": draw_ui_analytics(),
        "3.6": draw_ui_scanner(),
        "3.7": draw_ui_admin_dashboard(),
        "3.8": draw_ui_admin_table(),
        "appendix_structure": draw_project_structure(),
    }


def bibliography() -> list[str]:
    return [
        "1. Положение о выпускной квалификационной работе бакалавров, специалистов и магистров. Выпуск 3 [Электронный ресурс]. Локальный файл проекта образовательной организации (дата обращения: 14.04.2026).",
        "2. ГОСТ Р 7.0.100-2018. Система стандартов по информации, библиотечному и издательскому делу. Библиографическая запись. Библиографическое описание. Общие требования и правила составления. М.: Стандартинформ, 2018.",
        "3. ISO/IEC 18004:2015. Information technology. Automatic identification and data capture techniques. QR Code bar code symbology specification. Geneva: ISO, 2015.",
        "4. DENSO WAVE. QRcode.com [Электронный ресурс]. URL: https://www.qrcode.com/ (дата обращения: 14.04.2026).",
        "5. GS1. Introduction to 2D Barcodes [Электронный ресурс]. URL: https://www.gs1.org/standards/barcodes/2d-barcodes (дата обращения: 14.04.2026).",
        "6. Laravel. Release Notes [Электронный ресурс]. URL: https://laravel.com/docs/10.x/releases (дата обращения: 14.04.2026).",
        "7. Laravel. Routing [Электронный ресурс]. URL: https://laravel.com/docs/10.x/routing (дата обращения: 14.04.2026).",
        "8. Laravel. Eloquent: Relationships [Электронный ресурс]. URL: https://laravel.com/docs/10.x/eloquent-relationships (дата обращения: 14.04.2026).",
        "9. Laravel. Laravel Sanctum [Электронный ресурс]. URL: https://laravel.com/docs/10.x/sanctum (дата обращения: 14.04.2026).",
        "10. Vue.js. Guide [Электронный ресурс]. URL: https://vuejs.org/guide/introduction.html (дата обращения: 14.04.2026).",
        "11. Inertia.js. Documentation [Электронный ресурс]. URL: https://inertiajs.com/ (дата обращения: 14.04.2026).",
        "12. Vite. Guide [Электронный ресурс]. URL: https://vite.dev/guide/ (дата обращения: 14.04.2026).",
        "13. Tailwind CSS. Installation [Электронный ресурс]. URL: https://tailwindcss.com/docs/installation (дата обращения: 14.04.2026).",
        "14. Chart.js. Documentation [Электронный ресурс]. URL: https://www.chartjs.org/docs/latest/ (дата обращения: 14.04.2026).",
        "15. Fowler M. Patterns of Enterprise Application Architecture. Boston: Addison-Wesley, 2002.",
        "16. Sommerville I. Software Engineering. 10th ed. Boston: Pearson, 2015.",
        "17. MySQL 8.0 Reference Manual [Электронный ресурс]. URL: https://dev.mysql.com/doc/refman/8.0/en/ (дата обращения: 14.04.2026).",
        "18. Berners-Lee T., Fielding R., Masinter L. RFC 3986: Uniform Resource Identifier (URI): Generic Syntax [Электронный ресурс]. URL: https://www.rfc-editor.org/rfc/rfc3986 (дата обращения: 14.04.2026).",
        "19. OWASP Top 10:2021 [Электронный ресурс]. URL: https://owasp.org/www-project-top-ten/ (дата обращения: 14.04.2026).",
        "20. Docker Docs. Compose file reference [Электронный ресурс]. URL: https://docs.docker.com/reference/compose-file/ (дата обращения: 14.04.2026).",
        "21. README проекта qr-generator-fullstack-app [Электронный ресурс]. Локальный файл проекта: README.md (дата обращения: 14.04.2026).",
        "22. Маршруты и миграции проекта qr-generator-fullstack-app [Электронный ресурс]. Локальные файлы проекта: routes/web.php, routes/auth.php, routes/api.php, database/migrations/*.php (дата обращения: 14.04.2026).",
    ]


def body_content() -> dict[str, list[str] | tuple[list[str], list[list[str]]]]:
    return {
        "abstract": [
            "В выпускной квалификационной работе рассматривается разработка веб-сервиса для генерации и учета QR-кодов на базе стека Laravel 10, Vue 3, Inertia.js и сопутствующих библиотек проекта. Исследование выполнено по исходному коду приложения, конфигурации маршрутов, миграциям базы данных, frontend-страницам и административному модулю. Предметом анализа выступает не абстрактная модель генератора, а фактическая программная система, в которой уже реализованы генерация кода, сохранение изображений, динамические ссылки, учет сканирований и панель управления.",
            "В первой главе рассмотрены особенности предметной области сервисов QR-кодов, сопоставлены типовые подходы к статическим и динамическим кодам, выделены проблемы, связанные с хранением, редиректами, безопасностью и аналитикой. Во второй главе обоснован архитектурный подход, описаны технологии, структура репозитория, модель данных и взаимодействие модулей. В третьей главе подробно разобраны backend- и frontend-части, механизмы сканирования, административные сценарии, вопросы тестирования и дальнейшего развития проекта.",
            "Практическая значимость работы состоит в том, что описываемый сервис поддерживает полезный для бизнеса набор функций: генерацию кодов нескольких типов, динамические QR для тарифных планов Pro и Enterprise, фиксацию источников и параметров сканирования, отображение аналитики по отдельному коду, а также универсальный административный CRUD-интерфейс. Итоговый документ содержит новые диаграммы и иллюстрации, соответствующие тематике QR-сервиса, и оформлен по структуре и стилистике эталонной пояснительной записки.",
            "Ключевые слова: QR-код, динамический QR, веб-сервис, Laravel, Vue 3, Inertia.js, аналитика, сканирование, редирект, административная панель.",
        ],
        "introduction": [
            "Технология QR-кодов из вспомогательного инструмента маркировки давно превратилась в самостоятельный канал взаимодействия между цифровым сервисом и физическим объектом. В розничной торговле, логистике, мероприятиях, внутреннем документообороте и маркетинговых кампаниях QR-код одновременно выполняет функции идентификатора, ссылки, пропуска и носителя оперативной информации [1], [3]. Поэтому практическая ценность имеют не просто генераторы изображений, а сервисы, которые позволяют управлять жизненным циклом кода и анализировать его использование.",
            "Актуальность темы дополнительно усиливается распространением динамических QR-кодов. В отличие от статического варианта, динамический код содержит не конечный адрес, а промежуточную ссылку на серверное приложение, которое может перенаправить пользователя на новый ресурс, зафиксировать факт обращения и собрать телеметрию. Такой подход поддерживается стандартами отрасли и широко применяется в бизнес-процессах, где требуется гибкое управление ссылками и измерение эффективности кампаний [4], [5].",
            "Объектом исследования является веб-сервис генерации и учета QR-кодов как класс прикладных информационных систем. Предмет исследования составляют архитектурные, программные и организационные решения, необходимые для реализации генерации, хранения, редиректов, аналитики сканирований и администрирования пользователей. Работа ориентирована на реальный проект `qr-generator-fullstack-app`, размещенный в текущем рабочем пространстве и содержащий backend на Laravel, frontend на Vue и миграции базы данных.",
            "Цель работы заключается в разработке и обосновании архитектуры веб-сервиса, который объединяет генерацию QR-кодов нескольких типов, хранение результатов, динамические ссылки, учет сканирований, отображение статистики и административный контроль. Для достижения цели необходимо решить следующие задачи: проанализировать предметную область; определить функциональные и нефункциональные требования; спроектировать архитектуру и модель данных; описать реализацию основных модулей; подготовить диаграммы, иллюстрации и приложения по фактическому проекту.",
            "Практическая значимость исследования определяется возможностью непосредственного использования описываемого решения в учебном, коммерческом и внутрикорпоративном контуре. Сервис удобен там, где требуется быстро создавать QR-коды, отслеживать их сканирования и централизованно управлять контентом. Структура проекта и набор использованных технологий позволяют развивать систему без смены базовой архитектуры и без перехода к избыточно сложной микросервисной модели [2].",
        ],
        "1.1": [
            "Для современного сервиса QR-кодов недостаточно создать PNG-изображение по заданной строке. Пользователь ожидает, что система позволит выбрать тип данных, сохранить результат в личной истории, повторно открыть запись, скачать изображение и при необходимости заменить конечный адрес без перепечатки уже распространенного кода. Именно эта практическая логика отличает полноценный веб-сервис от локального генератора «в один клик».",
            "В рассматриваемом проекте такой подход выражен достаточно явно. Уже на уровне маршрутов видно разделение между публичными страницами, пользовательским контуром `/generate`, `/history`, `/qr/{id}/analytics` и административной зоной `/admin`. Модель `QrCode` хранит не только исходное содержимое, но и тип данных, JSON-payload, изображение, признак динамического режима и slug для редиректа. Модель `QrScan` фиксирует не только сам факт сканирования, но и страну, город, браузер, устройство и referer, что делает систему пригодной для базовой аналитики [6], [7].",
            "Особенность QR-сервиса состоит также в том, что он должен работать на стыке клиентских и серверных сценариев. Генерация и просмотр истории выполняются в защищенной пользовательской зоне, тогда как чтение динамического кода осуществляется публично и должно обрабатываться без авторизации. Это требует аккуратного разделения маршрутов, нормализации целевых URL и контроля владельца сущности при обращении к аналитике или удалению записей [8], [9].",
            "Следовательно, предметная область проекта связана не только с визуальным построением QR-кода, но и с сопровождением полного жизненного цикла сущности: от ввода параметров до сбора статистики по фактическому использованию. С инженерной точки зрения это уже не отдельная страница генератора, а прикладная информационная система с frontend, backend, моделью данных и административным управлением.",
        ],
        "1.2": [
            "При обзоре аналогов полезно различать два класса решений. Первый класс образуют легковесные генераторы, ориентированные на быстрое получение статического изображения. Они удобны для единичных задач, но обычно не предоставляют историю, группировку записей, доступ к статистике и разграничение ролей. Второй класс составляют маркетинговые платформы с динамическими QR-кодами, брендированием, глубокой аналитикой и командной работой.",
            "Исследуемый проект занимает промежуточную, но практичную позицию между этими крайностями. С одной стороны, приложение остается достаточно компактным, не требует отдельного SPA-API-слоя и может быть развернуто как единый Laravel-монолит. С другой стороны, в нем уже присутствуют те свойства, которые делают сервис полезным в реальной эксплуатации: управление тарифами, динамические ссылки, собственная страница аналитики и административный контур.",
            "По сравнению с минимальными веб-генераторами преимуществом текущей системы является то, что код можно не только создать, но и сопровождать: отфильтровать в истории, просмотреть число сканирований, проанализировать географию обращений и при необходимости удалить. По сравнению с крупными SaaS-платформами решение, разумеется, уступает в multi-tenant-возможностях, white-label-настройках и интеграциях с CRM, но выигрывает в прозрачности архитектуры и управляемости в рамках одной команды разработки [10], [11].",
            "Таким образом, проект логично рассматривать как учебно-прикладной сервис среднего уровня зрелости: он уже вышел за пределы демонстрационного генератора, но еще сохраняет архитектурную простоту, позволяющую подробно описывать каждое техническое решение в рамках ВКР.",
        ],
        "1.3": [
            "Первая проблема, которую должен решать подобный сервис, связана с разрывом между генерацией и дальнейшей эксплуатацией QR-кода. Если после создания изображения система не сохраняет запись и не поддерживает повторный доступ к ней, пользователь теряет возможность управлять ранее выпущенными кодами. В прикладной практике это приводит к дублированию, ошибкам распространения и потере связи между кодом и его назначением.",
            "Вторая проблема касается динамических QR-кодов. Когда один и тот же код используется в полиграфии, на упаковке или в инструкции, конечный адрес может меняться. Без серверного промежуточного слоя такие изменения потребуют выпуска нового изображения. Следовательно, системе требуется отдельный redirect-контур, который позволяет отделить визуальный идентификатор от конечного URI и при этом корректно обрабатывать запросы по публичному маршруту.",
            "Третья проблема — отсутствие объективной статистики по сканированиям. Для бизнеса важно знать не только количество созданных кодов, но и число обращений, используемые устройства, географию и источники переходов. Эти данные невозможно восстановить постфактум, если они не фиксируются в момент обращения пользователя к динамическому коду. Поэтому модуль учета сканирований должен проектироваться как базовая, а не второстепенная часть системы [12].",
            "Четвертая проблема связана с безопасностью и разграничением прав. QR-код часто содержит ссылку на внешний ресурс, а значит, приложение должно нормализовать целевой адрес, контролировать владельца аналитики, различать пользовательские и административные действия, защищать формы и не допускать несанкционированного доступа к чужим записям. Без решения этих задач сервис теряет практическую надежность.",
        ],
        "1.4": [
            "С учетом проведенного анализа задачу разработки можно сформулировать следующим образом: требуется создать веб-сервис генерации и учета QR-кодов, который объединяет пользовательский интерфейс создания кода, личную историю записей, поддержку динамических ссылок, учет сканирований и административное управление данными. При этом система должна оставаться понятной по архитектуре и воспроизводимой в развертывании.",
            "Цель работы в рамках ВКР состоит не только в перечислении функций, но и в обосновании того, почему выбранная архитектура действительно подходит для данного класса систем. Необходимо показать, как Laravel обеспечивает серверную логику и безопасность, каким образом Vue и Inertia используются для интерфейсов, как модель данных поддерживает генерацию и аналитику, и почему такой набор компонентов позволяет быстро развивать проект без неоправданного роста сложности [13].",
            "Практическая декомпозиция задачи включает несколько связанных подзадач. Нужно описать пользовательский сценарий генерации и сохранения кода, механизм публичного редиректа, хранение изображений и метаданных, страницу аналитики, административную панель и вспомогательные модули, связанные с тарифами и обратной связью. Кроме того, необходимо проанализировать структуру репозитория и показать, как она способствует сопровождаемости проекта.",
            "Таким образом, постановка задачи охватывает и предметную сторону сервиса, и инженерную реализацию: от функций, видимых пользователю, до маршрутов, таблиц базы данных, middleware, библиотек и правил развертывания.",
        ],
        "1.5_intro": [
            "Функциональные требования в настоящей работе выводятся не из внешнего рекламного описания, а из фактической кодовой базы проекта. Такой подход позволяет избежать вымышленных возможностей и зафиксировать именно тот набор сценариев, который подтверждается маршрутами, контроллерами, моделями, frontend-страницами и миграциями.",
            "На пользовательском уровне система должна обеспечивать регистрацию и вход, генерацию QR-кодов разных типов, просмотр истории, фильтрацию и сортировку записей, скачивание изображения, переход к аналитике динамических кодов, работу со страницей профиля, тарифами и формой обратной связи. Публичная часть должна включать страницу сканирования и маршрут редиректа `/r/{slug}`.",
            "На административном уровне сервис обязан предоставлять dashboard со сводными KPI, список управляемых таблиц, создание и редактирование записей, просмотр последних сканирований и распределения по странам. Отдельным функциональным блоком выступает автоматическое логирование параметров сканирования при обращении к динамическому QR-коду.",
        ],
        "1.5_table": (
            ["Группа требований", "Содержание"],
            [
                ["Генерация QR", "Выбор типа данных, настройка размера и цветов, формирование QR на клиенте и сохранение на сервере."],
                ["История записей", "Поиск по содержимому, фильтрация по статическому и динамическому режиму, сортировка и удаление QR-кодов."],
                ["Динамические коды", "Формирование slug, маршрут публичного редиректа, возможность изменения целевого контента без замены изображения."],
                ["Учет сканирований", "Фиксация даты, IP, страны, города, браузера, устройства и referer при обращении к динамическому коду."],
                ["Аналитика", "Отображение общего числа сканирований, списка событий и графика активности по дням."],
                ["Администрирование", "Dashboard, просмотр управляемых таблиц, создание, редактирование и удаление записей в административной зоне."],
            ],
        ),
        "1.6_intro": [
            "Нефункциональные требования для QR-сервиса не менее значимы, чем пользовательские функции. Система должна оставаться удобной в использовании, надежной при сохранении данных, безопасной при обработке публичных ссылок и сопровождаемой на уровне структуры кода. Именно эти свойства определяют, сможет ли приложение выйти за рамки демонстрационного прототипа [14], [15], [16].",
            "Требование производительности выражается в минимизации числа полных перезагрузок страниц и в отказе от избыточной распределенной архитектуры. Связка Laravel и Inertia позволяет быстро отдавать страницы и при этом сохранять единый серверный контур авторизации. Для истории QR-кодов применяются пагинация и фильтрация на стороне запроса, что ограничивает нагрузку на пользовательский интерфейс и базу данных.",
            "Требование надежности связано с тем, что QR-код является прикладной записью, а не одноразовой временной строкой. После сохранения пользователь должен иметь возможность повторно получить изображение, а динамический slug должен однозначно находить соответствующий объект. Учет сканирований также должен быть устойчивым к сбоям внешнего геосервиса: при недоступности API система продолжает работу, сохраняя запись с резервными значениями.",
            "Требование безопасности проявляется в необходимости разграничивать роли, защищать формы, валидировать входные данные и корректно обрабатывать внешние URI. Не менее важно требование сопровождаемости: код проекта должен быть организован так, чтобы контроллеры, страницы, модели и миграции были читаемы и не смешивались в одном слое.",
        ],
        "1.6_table": (
            ["Требование", "Реализация в проекте"],
            [
                ["Производительность", "Пагинация истории, клиентское превью QR, server-driven navigation через Inertia."],
                ["Надежность", "Сохранение записей и изображений в файловом хранилище, уникальный slug для dynamic QR, каскадные связи БД."],
                ["Безопасность", "Middleware `auth` и `admin`, проверка владельца аналитики, CSRF-защита, валидация запросов."],
                ["Масштабируемость", "Монолитная архитектура с четким разделением модулей и возможностью выделения аналитики в отдельный сервис."],
                ["Сопровождаемость", "Разделение на controllers, models, migrations, pages и layouts; предсказуемая структура репозитория."],
                ["Эксплуатационная пригодность", "Docker-конфигурация проекта, локальные seed-данные и административный интерфейс для контроля записей."],
            ],
        ),
        "2_intro": [
            "Во второй главе рассматриваются проектные решения, определяющие внутреннее устройство сервиса: архитектурный подход, стек технологий, структура репозитория, модель данных и взаимодействие модулей. Поскольку приложение уже реализовано в коде, проектирование в данной работе не является чисто умозрительным: каждая позиция подтверждается файлами проекта и прослеживается до конкретного маршрута, модели или frontend-страницы.",
        ],
        "2.1": [
            "При проектировании сервиса генерации и учета QR-кодов можно пойти по нескольким путям. На одном полюсе находится простое server-rendered-приложение с HTML-шаблонами и минимальным JavaScript. На другом — полностью разделенный стек с отдельным SPA-клиентом и API. Для рассматриваемой задачи выбран промежуточный вариант: Laravel-монолит с интерфейсами на Vue 3, доставляемыми через Inertia.js.",
            "Такой выбор оправдан спецификой проекта. Сервису нужны авторизация, работа с сессией, формы, защищенные маршруты, административная зона и несколько интерактивных экранов. Полноценное разделение на отдельный API и SPA увеличило бы стоимость разработки и сопровождения, но не дало бы принципиального выигрыша при текущем масштабе предметной области. В то же время чистый blade-подход усложнил бы интерфейсы генератора, сканера и аналитики.",
            "На уровне серверной логики монолитная архитектура с доменно-ориентированным разделением по каталогам остается наиболее прагматичным решением. Контроллеры открывают сценарии, модели описывают сущности, middleware формируют периметр доступа, а frontend-страницы остаются связаны с backend через Inertia-props. Такой подход соответствует паттернам корпоративных приложений, где важна централизованная координация доменной логики и инфраструктуры [15].",
            "Следовательно, архитектуру системы можно охарактеризовать как прикладной Laravel-монолит с клиентскими экранами на Vue 3, умеренным использованием JavaScript-библиотек и единым серверным контуром аутентификации и хранения данных.",
        ],
        "2.2": [
            "Базовой платформой backend-части выступает Laravel 10. Фреймворк обеспечивает маршрутизацию, middleware, Eloquent ORM, встроенную систему валидации и интеграцию с Sanctum. В контексте данного проекта особенно важны предсказуемость web-маршрутов и удобное описание отношений между сущностями `User`, `Plan`, `QrCode`, `QrScan` и `Feedback` [6], [7], [8].",
            "На стороне клиента используется Vue 3. Его роль состоит не в построении полностью автономного SPA, а в создании удобных интерактивных страниц внутри Inertia-приложения. Такой выбор позволяет использовать компонентный подход для генератора, истории, аналитики и административного dashboard, сохраняя при этом единый механизм передачи данных от Laravel к интерфейсу [10], [11].",
            "Для сборки frontend-ресурсов выбран Vite. Это современный инструмент, хорошо интегрированный с экосистемой Laravel и Vue, а Tailwind CSS используется как утилитарная база для части интерфейсных компонентов. Для визуализации графиков применяется Chart.js, для генерации изображений QR на сервере — пакет `simple-qrcode`, а для распознавания браузера и устройства при сканировании — библиотека `jenssegers/agent` [12], [13], [14].",
            "В качестве хранилища проект использует реляционную модель данных, совместимую с SQLite и MySQL. Такой выбор естественен для сервиса, где есть пользователи, тарифы, QR-коды и события сканирования, связанные отношениями один-ко-многим. Реляционная СУБД позволяет сохранять целостность, строить выборки по аналитике и поддерживать административный CRUD без дополнительного слоя хранилищ.",
        ],
        "2.2_table": (
            ["Технология", "Назначение в системе"],
            [
                ["Laravel 10", "HTTP-слой, маршрутизация, middleware, ORM Eloquent, базовая безопасность и серверная логика."],
                ["PHP 8.1+", "Исполнение backend-кода, обработка запросов и интеграция с пакетами проекта."],
                ["Vue 3", "Компонентный frontend для страниц генератора, истории, аналитики и административных экранов."],
                ["Inertia.js", "Связка серверного приложения и Vue-страниц без построения отдельного REST API."],
                ["Vite", "Сборка frontend-ресурсов и dev-сценарий с быстрым обновлением интерфейсов."],
                ["simple-qrcode", "Формирование PNG-изображений QR-кодов на серверной стороне."],
                ["Chart.js", "Построение графиков активности на странице аналитики QR-кода."],
                ["SQLite / MySQL", "Хранение пользователей, QR-кодов, сканирований, тарифов и обратной связи."],
            ],
        ),
        "2.3": [
            "Структура репозитория проекта организована последовательно и достаточно прозрачно. В каталоге `app/Http` размещены контроллеры и middleware, что позволяет сразу увидеть входные точки пользовательских и административных сценариев. Отдельного внимания заслуживает `QrCodeController`, в котором сосредоточены создание кода, выдача истории, обработка редиректа и построение аналитики. Наличие выделенного `AdminController` упрощает сопровождение административного контура.",
            "Модели предметной области находятся в `app/Models`. По их составу легко восстановить основные сущности проекта: пользователь, тариф, QR-код, сканирование и обратная связь. Миграции из каталога `database/migrations` подтверждают структуру таблиц, добавление полей для dynamic-режима, хранение payload и административного признака пользователя. Благодаря этому модель данных читается не только по тексту ВКР, но и по фактическому коду проекта [17].",
            "Frontend организован в `resources/js`. Здесь разделены layouts, переиспользуемые компоненты и страницы приложения. Подход удобен тем, что основные экраны — `QrGenerator.vue`, `QrHistory.vue`, `QrAnalytics.vue`, `QrScanner.vue`, `Admin/Index.vue` и `Admin/TableIndex.vue` — сосредоточены в одном месте и логически соответствуют пользовательским разделам сервиса. Это снижает стоимость навигации по проекту для разработчика и облегчает эволюцию интерфейса.",
            "Инфраструктурный слой размещен в верхнеуровневых файлах `composer.json`, `package.json`, `Dockerfile`, `docker-compose.yml`, а маршруты собраны в каталоге `routes`. В сумме структура проекта демонстрирует хорошую инженерную дисциплину: backend, frontend, схема БД и конфигурация развертывания не смешаны между собой и образуют понятный каркас приложения.",
        ],
        "2.4": [
            "Модель данных сервиса строится вокруг сущности `qr_codes`. В ней сохраняются владелец кода, тип, исходное содержимое, структурированный payload, путь к изображению, размер, цвета, признак динамического режима и slug. Тем самым один объект охватывает как визуальные параметры генерации, так и прикладные свойства, необходимые для последующего редиректа и аналитики. Такое решение сокращает число вспомогательных таблиц и оставляет структуру достаточно компактной.",
            "Сущность `qr_scans` выступает журналом фактов обращения к динамическому QR-коду. Каждая запись связана с конкретным кодом и содержит дату события, IP-адрес, страну, город, браузер, устройство и referer. Даже если часть сведений недоступна, сама схема остается полезной, поскольку фиксирует минимально необходимый набор параметров для построения графиков и списков активности. В дальнейшем на этой основе может быть выделена отдельная таблица агрегированной статистики по периодам.",
            "Сущность `users` дополняется ссылкой на тарифный план, а `plans` определяют уровень доступа к возможностям сервиса. В текущем проекте именно тариф влияет на доступность dynamic-режима. Такой механизм не перегружает доменную модель, но позволяет аккуратно ограничивать премиальные функции и использовать платные сценарии без отдельного billing-ядра. Дополнительная таблица `feedback` поддерживает канал связи пользователя с системой и может использоваться как операционный контур поддержки.",
            "С точки зрения дальнейшего масштабирования модель естественно расширяется сущностями `projects` и `qr_categories`, которые позволяют группировать коды по бизнес-направлениям, кампаниям или внутренним подразделениям. В текущей реализации эта роль частично покрывается фильтрацией по истории и типу кода, однако проектная ER-диаграмма показывает, в каком направлении схема может развиваться без ломки существующего ядра.",
        ],
        "2.5": [
            "Взаимодействие модулей в рассматриваемой системе строится вокруг последовательности «страница — маршрут — контроллер — модель/хранилище». Пользовательский сценарий начинается на Vue-странице: генератор собирает данные о типе QR-кода, содержимом, размере и цветах, затем отправляет запрос на backend. Контроллер валидирует данные, при необходимости формирует dynamic URL, создает PNG-изображение и сохраняет запись в таблице `qr_codes`.",
            "Сценарий чтения динамического кода устроен иначе. Он начинается вне защищенного пользовательского контура: браузер или мобильная камера открывают публичный маршрут `/r/{slug}`. Контроллер ищет соответствующий QR-код, при необходимости обращается к геосервису, определяет сведения об устройстве по `user-agent`, сохраняет событие в `qr_scans`, нормализует конечную цель и выполняет redirect. Тем самым один запрос одновременно завершает пользовательскую задачу и пополняет аналитический слой [18].",
            "Модуль аналитики повторно использует уже накопленные данные. При обращении к странице `/qr/{id}/analytics` backend загружает QR-код вместе со связанными сканированиями, считает агрегаты и передает результат в Inertia. На клиенте Chart.js строит график активности, а таблица событий показывает детализацию по странам, городам и устройствам. Такой способ разделения обязанностей позволяет держать вычислительную логику на сервере, а визуализацию — на стороне браузера.",
            "Административный модуль взаимодействует с теми же данными через универсальный CRUD-контур. `AdminController` динамически определяет структуру таблиц, поддерживает список управляемых сущностей и формирует dashboard со сводными показателями. В результате проект получает единый набор данных для пользовательского, аналитического и административного сценариев без дублирования бизнес-логики.",
        ],
        "2.6": [
            "Для сервиса, обрабатывающего публичные ссылки и пользовательские данные, вопросы безопасности и эксплуатации являются не второстепенными, а базовыми. Прежде всего приложение должно отличать защищенные пользовательские маршруты от публичного redirect-контура и административной зоны. В проекте это достигается комбинацией middleware `auth`, `verified` и `admin`, а также явной проверкой владельца сущности перед выдачей аналитики или изменением данных [19].",
            "Не менее важна корректная работа с URI. QR-код может ссылаться как на веб-страницу, так и на `mailto:`, `tel:` или `geo:`-схему, поэтому backend обязан различать допустимые цели и безопасно перенаправлять пользователя. Ошибочная обработка подобных значений способна привести либо к поломке сценария, либо к некорректному поведению браузера. По этой причине в проекте используются нормализация целевого адреса и серверная проверка динамического контента.",
            "С эксплуатационной точки зрения проект остается достаточно легким. Наличие `Dockerfile` и `docker-compose.yml` позволяет воспроизводить окружение, а локальные миграции и сиды дают разработчику быстрый старт. Для учебно-прикладного сервиса это оптимальный баланс: система уже готова к контейнеризации и автоматизированной сборке, но не требует избыточной платформенной сложности вроде отдельного service mesh или шины событий [20], [21], [22].",
        ],
        "3_intro": [
            "Третья глава посвящена фактической реализации проекта. Здесь рассматриваются контроллеры, модели, frontend-страницы и административные экраны, а также те сценарии, которые напрямую обеспечивают генерацию QR-кодов, публичные редиректы и учет аналитики. В отличие от проектной главы, акцент делается на том, как именно выбранные решения проявляются в кодовой базе и интерфейсе.",
        ],
        "3.1": [
            "Backend приложения построен вокруг маршрутов из `routes/web.php`, `routes/auth.php`, `routes/api.php` и контроллеров, отвечающих за разные контуры системы. Наиболее значимым является `QrCodeController`, так как в нем сосредоточены все ключевые операции предметной области: выдача списка QR-кодов пользователя, сохранение новых записей, удаление, обработка публичного slug и формирование страницы аналитики.",
            "Метод `store` иллюстрирует практическую логику сервиса. Контроллер проверяет права пользователя на создание динамического кода в зависимости от тарифного плана, валидирует входные параметры, при необходимости строит публичный URL для redirect-кода, а затем вызывает библиотеку `simple-qrcode` для генерации PNG-изображения. После этого запись сохраняется в базе вместе с метаданными, что обеспечивает связь между визуальным кодом и его серверным представлением.",
            "Метод `redirect` реализует публичный сценарий использования dynamic QR. Он ищет запись по slug, определяет IP-адрес клиента, извлекает данные браузера и устройства, пытается получить географические сведения через внешний сервис и сохраняет событие в `QrScan`. После этого контроллер нормализует цель и выполняет redirect либо возвращает текстовый ответ, если конечное содержимое не является внешней ссылкой. В результате redirect-контур одновременно решает прикладную задачу и пополняет аналитику.",
            "С инженерной точки зрения backend проекта выглядит достаточно зрелым для учебной работы: контроллеры отделены по ролям, модели компактны и понятны, а административный модуль использует общие возможности схемы БД без дублирования CRUD-кода. Это создает хороший фундамент для дальнейшего вынесения части логики в сервисный слой при росте требований.",
        ],
        "3.2": [
            "Frontend-представление проекта охватывает все основные пользовательские сценарии. Страница `Home.vue` выступает точкой входа и направляет пользователя к генерации, сканированию и истории. `QrGenerator.vue` реализует пошаговый ввод данных для нескольких типов QR-кодов, мгновенное превью на canvas, настройку размера и цветов и выбор динамического режима в зависимости от тарифа. Эта страница фактически является центральным рабочим экраном сервиса.",
            "Страница `QrHistory.vue` отвечает за сопровождение уже созданных записей. Здесь пользователь может искать коды по содержимому, фильтровать список по статическому и динамическому режиму, менять порядок сортировки, скачивать изображения, копировать исходное содержимое и переходить к аналитике. Такой набор действий показывает, что история выполняет роль не архива, а полноценного рабочего реестра QR-кодов.",
            "Страница `QrAnalytics.vue` строит график активности и выводит таблицу детальных событий. В ней соединяются данные, подготовленные backend, и визуализация на базе Chart.js. Отдельно выделяется `QrScanner.vue`, где используется браузерная камера и библиотека ZXing для считывания кода непосредственно в интерфейсе приложения. Это полезно как для демонстрации возможностей сервиса, так и для ускоренной проверки сформированных QR-кодов.",
            "Административный frontend вынесен в отдельные страницы `Admin/Index.vue`, `Admin/TableIndex.vue` и `Admin/TableForm.vue`. Такое решение поддерживает четкое различие между пользовательским и служебным интерфейсом, но при этом не разрывает единое визуальное и навигационное пространство приложения.",
        ],
        "3.3": [
            "Сценарий генерации и сохранения QR-кода объединяет клиентскую и серверную части системы. Сначала пользователь выбирает тип данных, вводит содержимое, задает размер и цветовую схему. Далее фронтенд формирует запрос, содержащий тип, человекочитаемое содержимое и вспомогательный payload. На сервере контроллер решает, будет ли код статическим или динамическим, а затем сохраняет изображение в `public/qr_codes` и создает запись в таблице `qr_codes`.",
            "Для динамического режима особенно важно, что в изображение не встраивается конечная ссылка напрямую. Вместо этого QR-код указывает на внутренний маршрут `/r/{slug}`. Такой подход позволяет впоследствии изменять целевое содержимое, не нарушая уже опубликованное изображение. Именно здесь соединяются требования предметной области: устойчивость ссылки, возможность аналитики и управляемость жизненным циклом кода.",
            "Работа с базой данных в данном сценарии достаточно прямолинейна, но функционально значима. Таблица `qr_codes` служит центром всей предметной модели: от нее зависят история пользователя, аналитика и административные операции. Дополнительные поля `type`, `payload`, `is_dynamic` и `slug` превращают таблицу из простого списка изображений в полноценный реестр QR-объектов с бизнес-смыслом.",
            "Таким образом, генерация и хранение в проекте реализованы корректно: создание кода, файловое сохранение, доступ к истории и публичное использование одного и того же объекта опираются на единый набор данных и не требуют дублирующих сущностей.",
        ],
        "3.4": [
            "Аналитический контур проекта строится на сочетании динамических QR-кодов и журнала сканирований. Пользователь видит лишь итоговую страницу аналитики, однако за ней стоит последовательность связанных действий: redirect-маршрут фиксирует событие, сохраняет параметры обращения, а затем отдельный экран агрегирует уже накопленные данные по конкретному коду. Такой подход дает системе измеряемость и позволяет оценивать эффективность отдельных QR-кампаний.",
            "Учет сканирований в текущем проекте включает дату, IP, страну, город, устройство, браузер и referer. Наличие этих полей достаточно для базовой статистики и для построения наглядных отчетов без сложного data-warehouse. Даже если геолокационный сервис временно недоступен, приложение продолжает работать и сохраняет запись, что важно для эксплуатационной устойчивости модуля.",
            "Управление типами QR-кодов решено на уровне поля `type` и соответствующей клиентской формы. Сервис поддерживает текстовые, ссылочные, почтовые, телефонные, SMS- и географические сценарии, а payload помогает восстановить прикладную структуру исходных данных. Для дальнейшего развития модель естественно расширяется выделением категорий QR-кодов и проектных группировок, но уже текущая реализация показывает жизнеспособность подхода и достаточную гибкость для реального применения.",
            "В результате аналитика, журнал сканирований и поддержка разных типов данных образуют связанный функциональный блок, который выгодно отличает рассматриваемый сервис от простых генераторов изображений.",
        ],
        "3.5": [
            "Безопасность проекта строится на нескольких уровнях. Первый уровень — аутентификация и проверка принадлежности маршрутов. Пользовательские страницы генерации, истории, профиля, тарифов и обратной связи доступны только после авторизации, а административная зона дополнительно защищена проверкой признака `is_admin`. Такой подход исключает доступ обычного пользователя к системным таблицам и служебным функциям.",
            "Второй уровень — защита данных и контроль владельца сущности. Аналитика по QR-коду доступна только владельцу записи, что прямо проверяется в контроллере. Аналогичным образом удаление или просмотр пользовательских данных не должны выполняться по чужим идентификаторам. Благодаря такому ограничению история QR-кодов остается персонализированным рабочим пространством, а не общим незащищенным реестром.",
            "Третий уровень — валидация и нормализация входных данных. При сохранении QR-кода сервер проверяет обязательность полей, размер, цвета и типы значений. Для публичного редиректа важно также корректно интерпретировать конечный URI и не допустить неконтролируемого поведения приложения. В совокупности эти меры соответствуют базовым рекомендациям secure-by-default для веб-приложений [19].",
            "Следовательно, безопасность в проекте реализована не как отдельный декларативный раздел, а как набор конкретных защитных механизмов, встроенных в маршруты, контроллеры и модель доступа.",
        ],
        "3.6": [
            "Административная панель проекта представляет собой не декоративный экран, а полноценный инструмент сопровождения сервиса. Dashboard показывает суммарное число сканирований, активность за текущий день, долю динамических QR-кодов, распределение по странам и последние события. Эти сведения помогают оператору быстро понять текущее состояние системы без обращения к базе данных напрямую.",
            "Особенно удачным решением является универсальный CRUD-контур `AdminController`. Вместо разработки отдельного контроллера для каждой сущности приложение использует динамическое определение структуры таблиц, формирует список управляемых таблиц и предоставляет экраны просмотра, создания, редактирования и удаления записей. Для учебного и внутреннего корпоративного проекта это практичный компромисс между скоростью разработки и функциональной полнотой.",
            "Административная зона также полезна как инженерный инструмент проверки корректности миграций и seed-данных. Поскольку оператор может открыть таблицы `users`, `plans`, `qr_codes`, `qr_scans` и `feedback`, структура данных становится наблюдаемой без использования отдельных BI-средств или SQL-клиента. Это делает приложение удобнее для поддержки и демонстрации.",
            "В целом административный модуль подтверждает, что проект ориентирован не только на конечного пользователя, но и на эксплуатацию: данные можно контролировать, агрегаты — анализировать, а записи — оперативно редактировать через веб-интерфейс.",
        ],
        "3.7": [
            "Тестовый контур проекта нельзя назвать исчерпывающим, однако в нем уже присутствуют важные практические элементы. Репозиторий содержит миграции и сиды, позволяющие быстро получить воспроизводимый набор данных: три тарифных плана, фиксированные пользовательские записи, десятки QR-кодов, события сканирования и обращения в поддержку. Это существенно упрощает ручную проверку интерфейсов и серверных сценариев.",
            "Проверка сборки для такого сервиса включает как минимум три уровня. Первый уровень — корректное выполнение миграций и формирование тестового наполнения. Второй — сборка frontend-ресурсов через Vite и визуальная проверка страниц генератора, истории, аналитики, сканера и dashboard. Третий — проход по ключевым пользовательским маршрутам: создание QR-кода, переход по динамической ссылке, просмотр аналитики и административное редактирование записей.",
            "Наличие Docker-конфигурации дополнительно повышает воспроизводимость результатов. Даже без сложного CI-контура проект уже пригоден для стандартного dev-процесса: разработчик может поднять окружение, применить миграции, загрузить сиды и проверить сценарии на подготовленных данных. Для ВКР это важное доказательство того, что система доведена до инженерно состоятельного состояния, а не существует только на уровне отдельных фрагментов кода.",
            "Перспективным шагом развития тестового слоя стало бы добавление feature- и browser-тестов для генерации dynamic QR, маршрута `/r/{slug}` и страницы аналитики. Однако уже текущий уровень организации проверки можно считать достаточной базой для дальнейшего роста.",
        ],
        "3.8": [
            "Перспективы развития проекта связаны прежде всего с углублением аналитики. В текущем виде система фиксирует события сканирования и показывает базовые графики, однако дальнейшее развитие может включать агрегацию по периодам, уникальным устройствам, UTM-меткам и проектным группировкам. Это позволит использовать сервис не только как генератор, но и как инструмент оценки эффективности кампаний и материалов.",
            "Второе направление — расширение доменной модели за счет явных сущностей `projects` и `qr_categories`. Такой шаг упростит управление большими наборами кодов, даст пользователю возможность группировать записи по задачам, а администратору — строить более точную отчетность. Архитектура текущего проекта допускает такое расширение без отказа от существующего ядра.",
            "Третье направление касается усиления безопасности и эксплуатационного контура. Сервис может быть дополнен rate limiting для redirect-маршрута, более строгими правилами валидации внешних URI, системой логирования инцидентов, оповещениями администратора и автоматическим резервным копированием файлов QR-изображений. Для production-развертывания также полезны централизованный secret management и отдельная очередь фоновых задач.",
            "Наконец, логичным развитием станет интеграция с внешними системами: CRM, корпоративными каталогами, сервисами рассылок и аналитическими платформами. При сохранении текущей архитектуры такие интеграции лучше вводить постепенно, не разрушая удачно выбранный баланс между простотой монолита и функциональной насыщенностью приложения.",
        ],
        "conclusion": [
            "В результате выполнения работы была разработана и описана система генерации и учета QR-кодов, построенная на Laravel 10, Vue 3 и Inertia.js. Анализ предметной области показал, что практическая ценность QR-сервиса определяется не столько самим фактом генерации изображения, сколько возможностью сопровождать код после его выпуска: хранить записи, управлять динамическими ссылками, фиксировать сканирования и анализировать накопленные данные.",
            "В первой главе были сформулированы требования к системе и выделены ключевые проблемы предметной области. Во второй главе обоснован выбор архитектуры, описан технологический стек, структура репозитория и модель данных. В третьей главе рассмотрена фактическая реализация backend- и frontend-контуров, публичного redirect-маршрута, аналитики и административной панели. Тем самым работа охватывает как концептуальный, так и прикладной уровень разработки.",
            "К важным результатам следует отнести подтверждение того, что даже сравнительно компактный Laravel-проект может быть организован как полноценный прикладной сервис, а не как набор разрозненных экранов. В системе уже присутствуют ключевые свойства инженерно зрелого решения: авторизация, разграничение ролей, учет событий, воспроизводимые данные, административный контроль и понятная структура кода.",
            "Практическая перспектива проекта связана с развитием проектных группировок, углублением аналитики, усилением тестового контура и постепенной интеграцией с внешними сервисами. Однако и в текущем состоянии описываемый веб-сервис может рассматриваться как рабочая основа для задач генерации и учета QR-кодов в учебной, коммерческой и внутрикорпоративной среде.",
        ],
    }


def main() -> None:
    assets = build_assets()
    content = body_content()
    doc = Document(str(TEMPLATE_PATH))
    set_update_fields(doc)

    cover_updates = {
        8: "«Разработка веб-сервиса для генерации и учета QR-кодов»",
        14: "Обучающийся: Бирюкова ______________________",
        23: "Тема ВКР: «Разработка веб-сервиса для генерации и учета QR-кодов».",
        24: "Исходные данные: исходный код проекта qr-generator-fullstack-app, маршруты приложения, миграции базы данных, frontend-страницы, административный модуль, Docker-конфигурация и seed-данные.",
        26: "1. Проанализировать предметную область использования QR-кодов и существующие аналоги.",
        27: "2. Сформулировать функциональные и нефункциональные требования к веб-сервису.",
        28: "3. Спроектировать архитектуру, модель данных и взаимодействие модулей системы.",
        29: "4. Описать реализацию backend, frontend, аналитики, безопасности и административной панели.",
        30: "5. Подготовить новые диаграммы, иллюстрации, приложения и список литературы по теме QR-сервиса.",
    }
    for idx, text in cover_updates.items():
        doc.paragraphs[idx].text = text
        for run in doc.paragraphs[idx].runs:
            set_run_font(run, 12)

    remove_body_after_prefix(doc, keep_body_index=40)

    page_break(doc)
    add_heading(doc, "РЕФЕРАТ", 1)
    for paragraph in content["abstract"]:
        add_text(doc, paragraph)

    page_break(doc)
    add_heading(doc, "ВВЕДЕНИЕ", 1)
    for paragraph in content["introduction"]:
        add_text(doc, paragraph)

    page_break(doc)
    add_heading(doc, "1 Анализ предметной области сервисов QR-кодов", 1)
    add_text(doc, "В первой главе исследуются особенности QR-сервисов, типовые проблемы сопровождения кодов и требования к системе на основе фактических возможностей проекта.", style="No Spacing")
    add_picture(doc, assets["1.1"], 15.8)
    add_caption(doc, "Рисунок 1.1 – Архитектурная схема веб-сервиса генерации и учета QR-кодов")

    add_heading(doc, "1.1 Особенности сервисов QR-кодов и контекст проекта", 2)
    for paragraph in content["1.1"]:
        add_text(doc, paragraph)

    add_heading(doc, "1.2 Обзор аналогов и практик применения", 2)
    for paragraph in content["1.2"]:
        add_text(doc, paragraph)

    add_heading(doc, "1.3 Проблемы, которые должна решать система", 2)
    for paragraph in content["1.3"]:
        add_text(doc, paragraph)

    add_heading(doc, "1.4 Постановка задачи и цели разработки", 2)
    for paragraph in content["1.4"]:
        add_text(doc, paragraph)

    add_heading(doc, "1.5 Функциональные требования к сервису", 2)
    for paragraph in content["1.5_intro"]:
        add_text(doc, paragraph)
    add_caption(doc, "Таблица 1.1 – Основные функциональные требования")
    headers, rows = content["1.5_table"]  # type: ignore[assignment]
    add_table(doc, headers, rows, widths=[5.0, 11.0])

    add_heading(doc, "1.6 Нефункциональные требования", 2)
    for paragraph in content["1.6_intro"]:
        add_text(doc, paragraph)
    add_caption(doc, "Таблица 1.2 – Ключевые нефункциональные требования")
    headers, rows = content["1.6_table"]  # type: ignore[assignment]
    add_table(doc, headers, rows, widths=[5.0, 11.0])

    page_break(doc)
    add_heading(doc, "2 Проектирование системы", 1)
    for paragraph in content["2_intro"]:
        add_text(doc, paragraph, style="No Spacing")
    add_picture(doc, assets["2.1"], 16.2)
    add_caption(doc, "Рисунок 2.1 – ER-диаграмма модели данных проекта")
    add_picture(doc, assets["2.2"], 16.0)
    add_caption(doc, "Рисунок 2.2 – Диаграмма вариантов использования")

    add_heading(doc, "2.1 Выбор архитектурного подхода", 2)
    for paragraph in content["2.1"]:
        add_text(doc, paragraph)

    add_heading(doc, "2.2 Выбор технологий и зависимостей", 2)
    for paragraph in content["2.2"]:
        add_text(doc, paragraph)
    add_caption(doc, "Таблица 2.1 – Основные технологии проекта")
    headers, rows = content["2.2_table"]  # type: ignore[assignment]
    add_table(doc, headers, rows, widths=[4.0, 12.0])

    add_heading(doc, "2.3 Структура проекта и распределение ответственности между каталогами", 2)
    for paragraph in content["2.3"]:
        add_text(doc, paragraph)
    add_picture(doc, assets["appendix_structure"], 15.5)
    add_caption(doc, "Рисунок 2.3 – Структура каталогов проекта")

    add_heading(doc, "2.4 Проектирование модели данных и ER-диаграмма", 2)
    for paragraph in content["2.4"]:
        add_text(doc, paragraph)

    add_heading(doc, "2.5 Взаимодействие модулей", 2)
    for paragraph in content["2.5"]:
        add_text(doc, paragraph)

    add_heading(doc, "2.6 Безопасность и эксплуатационное проектирование", 2)
    for paragraph in content["2.6"]:
        add_text(doc, paragraph)

    page_break(doc)
    add_heading(doc, "3 Реализация веб-сервиса", 1)
    for paragraph in content["3_intro"]:
        add_text(doc, paragraph, style="No Spacing")
    add_picture(doc, assets["3.1"], 16.1)
    add_caption(doc, "Рисунок 3.1 – Последовательность создания динамического QR-кода и его сканирования")
    add_picture(doc, assets["3.2"], 16.1)
    add_caption(doc, "Рисунок 3.2 – Компонентная структура приложения")
    add_picture(doc, assets["3.3"], 15.5)
    add_caption(doc, "Рисунок 3.3 – Интерфейс генерации QR-кода")
    add_picture(doc, assets["3.4"], 15.5)
    add_caption(doc, "Рисунок 3.4 – Страница истории QR-кодов")
    add_picture(doc, assets["3.5"], 15.5)
    add_caption(doc, "Рисунок 3.5 – Страница аналитики QR-кода")
    add_picture(doc, assets["3.6"], 15.5)
    add_caption(doc, "Рисунок 3.6 – Страница сканирования QR-кодов")

    add_heading(doc, "3.1 Реализация backend-части", 2)
    for paragraph in content["3.1"]:
        add_text(doc, paragraph)

    add_heading(doc, "3.2 Реализация frontend: генератор, история, аналитика и сканер", 2)
    for paragraph in content["3.2"]:
        add_text(doc, paragraph)

    add_heading(doc, "3.3 Генерация, редиректы и работа с базой данных", 2)
    for paragraph in content["3.3"]:
        add_text(doc, paragraph)

    add_heading(doc, "3.4 Учет сканирований, аналитика и управление типами QR", 2)
    for paragraph in content["3.4"]:
        add_text(doc, paragraph)

    add_heading(doc, "3.5 Безопасность, авторизация и разграничение доступа", 2)
    for paragraph in content["3.5"]:
        add_text(doc, paragraph)

    add_heading(doc, "3.6 Административная панель и статистика", 2)
    for paragraph in content["3.6"]:
        add_text(doc, paragraph)
    add_picture(doc, assets["3.7"], 15.5)
    add_caption(doc, "Рисунок 3.7 – Административная панель: dashboard и ключевые метрики")
    add_picture(doc, assets["3.8"], 15.5)
    add_caption(doc, "Рисунок 3.8 – Административная панель: управление записями QR-кодов")

    add_heading(doc, "3.7 Тестирование, проверка сборки и верификация сценариев", 2)
    for paragraph in content["3.7"]:
        add_text(doc, paragraph)

    add_heading(doc, "3.8 Перспективы развития проекта", 2)
    for paragraph in content["3.8"]:
        add_text(doc, paragraph)

    page_break(doc)
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1)
    for paragraph in content["conclusion"]:
        add_text(doc, paragraph)

    page_break(doc)
    add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ", 1)
    for item in bibliography():
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        set_run_font(run, 12)

    page_break(doc)
    add_heading(doc, "ПРИЛОЖЕНИЯ", 1)
    add_text(doc, "В приложения вынесены материалы, которые дополняют основную часть и позволяют проверить состав маршрутов, структуру сущностей, зависимости проекта и эксплуатационные сценарии. Они не дублируют основные рисунки и таблицы главы 1–3.", style="Normal")

    routes = extract_routes()
    page_break(doc)
    add_heading(doc, "Приложение А", 1)
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Реестр маршрутов системы")
    add_text(doc, "Ниже приведен укрупненный реестр маршрутов, восстановленный по файлам `routes/web.php`, `routes/auth.php`, `routes/api.php` и `routes/health.php`. В таблицах сохранены HTTP-метод, URI, имя маршрута и базовый обработчик.")
    for idx, chunk in enumerate(chunked(routes, 10), start=1):
        start = (idx - 1) * 10 + 1
        end = start + len(chunk) - 1
        p = doc.add_paragraph(style="No Spacing")
        p.add_run(f"Фрагмент реестра маршрутов ({start}–{end})")
        add_table(doc, ["Метод", "URI", "Имя маршрута", "Контроллер / действие"], chunk, widths=[2.0, 4.0, 4.0, 6.0])

    page_break(doc)
    add_heading(doc, "Приложение Б", 1)
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Словарь сущностей базы данных")
    add_text(doc, "Таблица подготовлена по миграциям и моделям Eloquent. Она отражает ключевые сущности, задействованные в текущем состоянии проекта и в целевой модели развития QR-сервиса.")
    entities = [
        ["users", "Учетные записи пользователей и администраторов", "plan_id, email, password, is_admin"],
        ["plans", "Тарифные планы сервиса", "name, price, is_default"],
        ["qr_codes", "Основной реестр QR-кодов", "user_id, type, content, payload, slug, is_dynamic"],
        ["qr_scans", "Журнал событий сканирования", "qr_code_id, country, city, browser, device, referer"],
        ["feedback", "Обращения пользователей", "user_id, subject, category, status, priority"],
        ["personal_access_tokens", "Токены API и пользовательских сессий", "tokenable_id, token, abilities"],
        ["projects", "Проектная группировка QR-кодов", "user_id, name, description, is_archived"],
        ["qr_categories", "Категории QR-кодов", "project_id, name, slug, sort_order"],
        ["scan_statistics", "Агрегаты по периодам", "qr_code_id, period_date, scans_total"],
        ["redirect_rules", "Правила маршрутизации dynamic QR", "qr_code_id, target_url, is_active"],
    ]
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Фрагмент словаря сущностей (1–5)")
    add_table(doc, ["Таблица", "Назначение", "Ключевые поля"], entities[:5], widths=[3.0, 7.0, 6.0])
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Фрагмент словаря сущностей (6–10)")
    add_table(doc, ["Таблица", "Назначение", "Ключевые поля"], entities[5:], widths=[3.0, 7.0, 6.0])

    page_break(doc)
    add_heading(doc, "Приложение В", 1)
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Инвентарь модулей и зависимостей проекта")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Ключевые backend-модули")
    backend_modules = [
        ["QrCodeController", "Генерация, история, redirect и аналитика QR-кодов."],
        ["AdminController", "Dashboard и универсальный CRUD по таблицам."],
        ["PlanController", "Просмотр тарифов, подписка и имитация оплаты."],
        ["ProfileController", "Просмотр и обновление профиля пользователя."],
        ["FeedbackController", "Создание и просмотр обращений пользователя."],
        ["AdminMiddleware", "Разграничение доступа к административной зоне."],
    ]
    add_table(doc, ["Модуль", "Назначение"], backend_modules, widths=[5.0, 11.0])
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Backend-зависимости из composer.json")
    backend_deps = [
        ["laravel/framework", "Базовый фреймворк приложения."],
        ["inertiajs/inertia-laravel", "Связка Laravel и Inertia.js."],
        ["laravel/sanctum", "Маршруты и механизмы аутентификации токенов."],
        ["simplesoftwareio/simple-qrcode", "Генерация серверных QR-изображений."],
        ["jenssegers/agent", "Определение браузера и устройства по user-agent."],
    ]
    add_table(doc, ["Пакет", "Назначение"], backend_deps, widths=[6.0, 10.0])
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Frontend-зависимости из package.json")
    frontend_deps = [
        ["vue", "Компонентный frontend-фреймворк."],
        ["@inertiajs/vue3", "Интеграция Inertia и Vue."],
        ["vite", "Сборка frontend-ресурсов."],
        ["chart.js", "Графики на странице аналитики."],
        ["@zxing/library", "Сканирование QR-кодов через камеру."],
        ["qrcode", "Клиентское превью кода на canvas."],
    ]
    add_table(doc, ["Пакет", "Назначение"], frontend_deps, widths=[6.0, 10.0])

    page_break(doc)
    add_heading(doc, "Приложение Г", 1)
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Эксплуатационные сценарии и срез тестовых данных")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Основные эксплуатационные сценарии")
    ops = [
        ["Локальная разработка", "Установка зависимостей, запуск Vite и Laravel-окружения, миграции и сидирование."],
        ["Генерация QR", "Пользователь создает код, сохраняет PNG и получает запись в истории."],
        ["Публичный redirect", "Маршрут `/r/{slug}` перенаправляет клиента и создает запись в `qr_scans`."],
        ["Аналитика", "Пользователь открывает страницу QR-аналитики и получает график + список событий."],
        ["Администрирование", "Оператор просматривает dashboard и управляет таблицами через CRUD."],
        ["Обратная связь", "Пользователь отправляет сообщение, система назначает приоритет по тарифу."],
    ]
    add_table(doc, ["Сценарий", "Содержание"], ops, widths=[5.0, 11.0])
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Срез наполнения тестовой БД по сидерам проекта")
    dataset = [
        ["plans", "3"],
        ["users", "60"],
        ["qr_codes", "90"],
        ["qr_scans", "100"],
        ["feedback", "70"],
        ["preset QR images", "90 файлов seed-изображений"],
    ]
    add_table(doc, ["Сущность", "Количество записей"], dataset, widths=[8.0, 8.0])

    page_break(doc)
    add_heading(doc, "Приложение Д", 1)
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("Отчет системы «Антиплагиат»")
    add_text(doc, "Отчет системы «Антиплагиат» прикладывается к окончательной версии ВКР после проверки готового текста в корпоративной системе образовательной организации. В данной версии документа место под официальный отчет зарезервировано, но сам отчет не формируется автоматически средствами проекта.", style="Normal")

    doc.save(str(OUTPUT_PATH))
    print(json.dumps({"output": str(OUTPUT_PATH), "assets_dir": str(ASSETS_DIR)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
